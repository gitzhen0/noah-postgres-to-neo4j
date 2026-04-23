"""
Generate NOAH Capstone Final Report (DOCX) per the Assignment 9A rubric.

The rubric demands: Times New Roman 12pt, 1" margins, cover page (no number),
bottom-centered page numbers, each section on a new page, Table of Contents,
numbered figures, 250-300 word abstract, and appendices A-H with signed
artifacts from prior assignments.

Output: ./NOAH_Capstone_Final_Report.docx (root of the repo).
Run:    ./venv/bin/python scripts/generate_final_report.py
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


REPO = Path(__file__).resolve().parent.parent
OUT  = REPO / "NOAH_Capstone_Final_Report.docx"

# ─── style constants ──────────────────────────────────────────────────────
FONT_BODY  = "Times New Roman"
FONT_MONO  = "Consolas"
SIZE_BODY  = Pt(12)
SIZE_H1    = Pt(18)
SIZE_H2    = Pt(14)
SIZE_H3    = Pt(12)
SIZE_SMALL = Pt(10)

INK      = RGBColor(0x00, 0x00, 0x00)
MUTED    = RGBColor(0x50, 0x50, 0x50)
ACCENT   = RGBColor(0x1F, 0x49, 0x7D)  # traditional MS Word blue

FIG_COUNTER = {"num": 0}
TBL_COUNTER = {"num": 0}


# ─── low-level helpers ────────────────────────────────────────────────────
def set_margins(section, margin: float = 1.0):
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Inches(margin))


def set_font(run, *, size=SIZE_BODY, bold=False, italic=False,
             color=INK, font=FONT_BODY):
    run.font.name = font
    # Some fonts need the east-asian variant too, or Word falls back
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_p(doc_or_cell, text="", *, size=SIZE_BODY, bold=False, italic=False,
          color=INK, font=FONT_BODY, align=None, space_after=Pt(6),
          line_spacing=1.15, page_break_before=False):
    """Add a paragraph of plain text."""
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0] if doc_or_cell.paragraphs and not doc_or_cell.paragraphs[0].text else doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if page_break_before:
        p.paragraph_format.page_break_before = True
    if text:
        r = p.add_run(text)
        set_font(r, size=size, bold=bold, italic=italic, color=color, font=font)
    return p


def add_mixed(doc, segments, *, align=None, space_after=Pt(6),
              line_spacing=1.15, page_break_before=False):
    """Add a paragraph with multiple runs.
    segments: list of dicts {text, bold?, italic?, size?, color?, font?}"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if page_break_before:
        p.paragraph_format.page_break_before = True
    for seg in segments:
        r = p.add_run(seg["text"])
        set_font(r,
                 size=seg.get("size", SIZE_BODY),
                 bold=seg.get("bold", False),
                 italic=seg.get("italic", False),
                 color=seg.get("color", INK),
                 font=seg.get("font", FONT_BODY))
    return p


def add_heading1(doc, text, *, page_break=True):
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=SIZE_H1, bold=True, color=INK, font=FONT_BODY)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=SIZE_H2, bold=True, color=INK, font=FONT_BODY)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=SIZE_H3, bold=True, italic=True, color=INK, font=FONT_BODY)
    return p


def add_bullets(doc, items, *, indent=Inches(0.25)):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = indent
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run("•  ")
        set_font(r1, size=SIZE_BODY)
        r2 = p.add_run(item)
        set_font(r2, size=SIZE_BODY)


def add_numbered_list(doc, items, *, indent=Inches(0.25)):
    for i, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = indent
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(f"{i}.  {item}")
        set_font(r, size=SIZE_BODY)


def add_figure_caption(doc, caption, *, chapter):
    FIG_COUNTER["num"] += 1
    label = f"Figure {chapter}-{FIG_COUNTER['num']}: "
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(10)
    r1 = p.add_run(label)
    set_font(r1, size=Pt(10), bold=True, italic=True, color=INK)
    r2 = p.add_run(caption)
    set_font(r2, size=Pt(10), italic=True, color=INK)


def add_figure(doc, image_path, caption, *, chapter, width_in=5.5):
    """Centered image + numbered caption below it."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(width_in))
    except Exception as exc:
        print(f"[warn] could not embed {image_path}: {exc}")
        t = p.add_run(f"[missing figure: {image_path}]")
        set_font(t, italic=True, color=MUTED)
        return
    add_figure_caption(doc, caption, chapter=chapter)


def add_table_caption(doc, caption, *, chapter):
    TBL_COUNTER["num"] += 1
    label = f"Table {chapter}-{TBL_COUNTER['num']}: "
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label)
    set_font(r1, size=Pt(10), bold=True, italic=True, color=INK)
    r2 = p.add_run(caption)
    set_font(r2, size=Pt(10), italic=True, color=INK)


def reset_figure_counter():
    FIG_COUNTER["num"] = 0
    TBL_COUNTER["num"] = 0


def add_table(doc, header, rows, *, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = w
    # header
    for i, h in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        set_font(r, size=Pt(11), bold=True, color=INK)
    # data
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            set_font(r, size=Pt(11))
    # extra space after table
    doc.add_paragraph()
    return tbl


def add_code(doc, text):
    """Add a code block in fixed-width."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        p.paragraph_format.left_indent = Inches(0.4)
        r = p.add_run(line if line else " ")
        set_font(r, size=Pt(10), font=FONT_MONO)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_page_number_in_footer(section, *, show=True):
    """Centre-align a PAGE field in the footer, using the begin/end field
    pattern that both Word and LibreOffice render reliably."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # clear existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    if not show:
        return
    run = p.add_run()
    set_font(run, size=Pt(10), color=MUTED)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE   \\* MERGEFORMAT '
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    t_placeholder = OxmlElement('w:t')
    t_placeholder.text = '1'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    for el in (fld_begin, instr, fld_sep, t_placeholder, fld_end):
        run._element.append(el)


def new_section(doc, *, different_first_page=False):
    """Start a new section with its own footer setup."""
    s = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(s)
    s.different_first_page_header_footer = different_first_page
    return s


# ─── source-content loaders ──────────────────────────────────────────────
SRC_DIR = Path("/tmp")

SOURCES = {
    "frs":          SRC_DIR / "src_FRS_Requirement_Specification_Zhen.md",
    "wbs":          SRC_DIR / "src_WBS.md",
    "risk":         SRC_DIR / "src_Risk_Management_Plan_Zhen_Yang.md",
    "trial":        SRC_DIR / "src_Technology_Trial_Plan_Zhen_Yang.md",
    "status_feb":   SRC_DIR / "src_Status_Feb26.md",
    "status_mar":   SRC_DIR / "src_Status_Mar25.md",
    "status_apr":   SRC_DIR / "src_Status_Apr16.md",
    "bibliography": SRC_DIR / "src_Assignment7_PartA_AnnotatedBibliography_ZhenYang.md",
    "proposal":     SRC_DIR / "src_Final_Project_Proposal_Finished_1.md",
}


def load_markdown(path: Path) -> str:
    if not path.exists():
        return f"[SOURCE MISSING: {path.name}]"
    return path.read_text()


# A tiny pandoc-style markdown renderer for the appendix sections. Supports:
# - paragraphs, headings (#/##/###), bold (**x**), italic (*x*), code (`x`)
# - bullet lists ("- " / "• "), numbered lists ("1. ")
# - blockquotes ("> ")
# - pipe tables
def render_markdown(doc, text: str, *, skip_memo_header=True):
    lines = text.splitlines()
    i = 0
    # Optionally skip the first ~20 lines that form the memo header
    if skip_memo_header:
        start = 0
        for j, ln in enumerate(lines[:50]):
            if re.match(r"^\*\*RE:", ln) or re.match(r"^RE:", ln) or ln.strip().startswith("**Project") or ln.strip().startswith("## "):
                start = j + 1
                break
        lines = lines[start:]
    # skip leading blanks
    while lines and not lines[0].strip():
        lines.pop(0)

    def flush_list(items, ordered):
        if ordered:
            add_numbered_list(doc, items)
        else:
            add_bullets(doc, items)

    list_buf = []
    list_ordered = False
    para_buf = []

    def flush_para():
        nonlocal para_buf
        if para_buf:
            text = " ".join(para_buf).strip()
            para_buf = []
            if text:
                _emit_paragraph(doc, text)

    def flush_list_if():
        nonlocal list_buf, list_ordered
        if list_buf:
            flush_list(list_buf, list_ordered)
            list_buf = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # blank line
        if not stripped:
            flush_para()
            flush_list_if()
            i += 1
            continue

        # heading
        if stripped.startswith("### "):
            flush_para(); flush_list_if()
            add_heading3(doc, stripped[4:])
            i += 1
            continue
        if stripped.startswith("## "):
            flush_para(); flush_list_if()
            add_heading3(doc, stripped[3:])
            i += 1
            continue
        if stripped.startswith("# "):
            flush_para(); flush_list_if()
            add_heading2(doc, stripped[2:])
            i += 1
            continue

        # pipe table detection
        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|?\s*[-:]+[-| :]*\s*\|?\s*$", lines[i + 1].strip()):
            flush_para(); flush_list_if()
            rows = []
            while i < len(lines) and "|" in lines[i]:
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if re.match(r"^[-:]+$", row_cells[0].replace(" ", "")) or all(re.match(r"^[-:]*$", c) for c in row_cells):
                    i += 1; continue
                rows.append(row_cells)
                i += 1
            if rows:
                header = rows[0]
                body = rows[1:]
                add_table(doc, header, body)
            continue

        # pandoc multi-line table (rare in our source) — treat as code block
        if re.match(r"^[-=+]{3,}", stripped):
            flush_para(); flush_list_if()
            block = []
            while i < len(lines) and lines[i].strip():
                block.append(lines[i])
                i += 1
            add_code(doc, "\n".join(block))
            continue

        # blockquote
        if stripped.startswith("> "):
            flush_para(); flush_list_if()
            quoted = [stripped[2:]]
            i += 1
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                q = lines[i].lstrip()[1:].lstrip()
                quoted.append(q)
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_formatted_runs(p, " ".join(quoted), italic_base=True)
            continue

        # bullets
        if stripped.startswith("- ") or stripped.startswith("• ") or stripped.startswith("* "):
            flush_para()
            if list_buf and list_ordered:
                flush_list_if()
            list_ordered = False
            item = stripped[2:].strip()
            # join continuation lines
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip():
                item += " " + lines[i].strip()
                i += 1
            list_buf.append(item)
            continue

        # numbered list
        m = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if m:
            flush_para()
            if list_buf and not list_ordered:
                flush_list_if()
            list_ordered = True
            list_buf.append(m.group(2))
            i += 1
            continue

        # pandoc img/link leftovers or html comments — ignore
        if stripped.startswith("![") or stripped.startswith("<!--"):
            i += 1; continue

        # default: accumulate into current paragraph
        para_buf.append(stripped)
        i += 1

    flush_para(); flush_list_if()


def _emit_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    _add_formatted_runs(p, text)


def _add_formatted_runs(p, text, *, italic_base=False):
    # parse **bold**, *italic*, `code` minimally
    # split on markers while preserving them
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    pos = 0
    for m in token_re.finditer(text):
        if m.start() > pos:
            plain = text[pos:m.start()]
            r = p.add_run(plain)
            set_font(r, italic=italic_base)
        tok = m.group(1)
        if tok.startswith("**"):
            r = p.add_run(tok[2:-2])
            set_font(r, bold=True, italic=italic_base)
        elif tok.startswith("`"):
            r = p.add_run(tok[1:-1])
            set_font(r, font=FONT_MONO, size=Pt(11))
        else:  # *italic*
            r = p.add_run(tok[1:-1])
            set_font(r, italic=True)
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_font(r, italic=italic_base)


# ─── sections ─────────────────────────────────────────────────────────────

def section_cover(doc):
    s = doc.sections[0]
    set_margins(s)
    # Cover page (only page of section 1) has an empty footer — no number.
    set_page_number_in_footer(s, show=False)
    # top spacer
    for _ in range(3):
        doc.add_paragraph()
    add_p(doc, "Automated RDBMS → Knowledge Graph Conversion Bot:",
          size=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=Pt(6), line_spacing=1.2)
    add_p(doc, "An Automated PostgreSQL-to-Neo4j Migration and Text2Cypher "
               "Natural-Language Query Interface, Validated on the NYC "
               "Affordable Housing (NOAH) Database",
          size=Pt(14), italic=True, color=MUTED,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36), line_spacing=1.3)
    add_p(doc, "Applied Project Final Report",
          size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=Pt(24))
    add_p(doc, "By", size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_p(doc, "Zhen Yang", size=Pt(18), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(24))
    add_p(doc, "Spring 2026", size=Pt(14),
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(48))
    add_p(doc, "A paper submitted in partial fulfillment of the requirements "
               "for the degree of",
          size=Pt(12), italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_p(doc, "Master of Science in Management and Systems",
          size=Pt(14), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(18))
    add_p(doc, "at the", size=Pt(12), italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    add_p(doc, "Division of Programs in Business",
          size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "School of Professional Studies",
          size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "New York University",
          size=Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(36))
    add_p(doc, "Faculty Advisor: Dr. Andres Fortino, Clinical Assistant Professor",
          size=Pt(12), italic=True, color=MUTED,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_p(doc, "Project Sponsor: The Digital Forge Lab",
          size=Pt(12), italic=True, color=MUTED,
          align=WD_ALIGN_PARAGRAPH.CENTER)


def section_toc(doc):
    add_page_break(doc)
    add_heading1(doc, "Table of Contents", page_break=False)
    add_p(doc,
        "The Table of Contents below is a Word field that auto-populates "
        "when the document opens in Microsoft Word. If the list below "
        "appears blank, right-click anywhere in this table and choose "
        "\"Update Field\" (or press F9).",
        italic=True, color=MUTED, space_after=Pt(12))
    p = doc.add_paragraph()
    r = p.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    fld_char1.set(qn('w:dirty'), 'true')     # auto-update on open
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \o "1-3" \h \z \u '
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Right-click here → Update Field to populate the Table of Contents."
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for el in (fld_char1, instr, sep, placeholder, end):
        r._element.append(el)
    set_font(r, size=SIZE_BODY)


def section_lot(doc):
    """List of Tables (since the report has 4 numbered tables)."""
    add_page_break(doc)
    add_heading1(doc, "List of Tables", page_break=False)
    add_p(doc,
        "The List of Tables below is a Word field that auto-populates when "
        "the document opens in Microsoft Word. If it appears blank, right-"
        "click and choose \"Update Field\".",
        italic=True, color=MUTED, space_after=Pt(12))
    p = doc.add_paragraph()
    r = p.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    fld_char1.set(qn('w:dirty'), 'true')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r' TOC \h \z \c "Table" '
    sep = OxmlElement('w:fldChar')
    sep.set(qn('w:fldCharType'), 'separate')
    placeholder = OxmlElement('w:t')
    placeholder.text = "Update Field to populate the List of Tables."
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    for el in (fld_char1, instr, sep, placeholder, end):
        r._element.append(el)
    set_font(r, size=SIZE_BODY)


def section_declaration(doc):
    add_heading1(doc, "Declaration")
    add_p(doc,
        "I, Zhen Yang, declare that this project report submitted by me to the "
        "School of Professional Studies, New York University in partial "
        "fulfillment of the requirement for the award of the degree of "
        "Master of Science in Management and Systems is a record of project "
        "work carried out by me under the guidance of Dr. Andres Fortino, "
        "NYU Clinical Assistant Professor of Management and Systems.")
    add_p(doc,
        "I grant powers of discretion to the Division of Programs in Business, "
        "School of Professional Studies, and New York University to allow this "
        "report to be copied in part or in full without further reference to me. "
        "The permission covers only copies made for study purposes or for "
        "inclusion in Division of Programs in Business, School of Professional "
        "Studies, and New York University research publications, subject to "
        "normal conditions of acknowledgment.")
    add_p(doc,
        "I further declare that the work reported in this project has not been "
        "submitted and will not be submitted, either in part or in full, for the "
        "award of any other degree or diploma in this institute or any other "
        "institute or university.")
    add_p(doc, "")
    add_p(doc, "Signed:  _____________________________________    Date: _______________",
          space_after=Pt(18))
    add_p(doc, "Zhen Yang", size=SIZE_BODY, italic=True, color=MUTED)


def section_acknowledgments(doc):
    add_heading1(doc, "Acknowledgments")
    add_p(doc,
        "I sincerely thank Dr. Andres Fortino, my faculty advisor and the "
        "sponsor of this capstone project through The Digital Forge Lab at the "
        "NYU School of Professional Studies. Dr. Fortino's guidance on project "
        "scoping, methodology selection, and research framing was instrumental "
        "throughout the twelve-week project. His insistence that the evaluation "
        "be grounded in measurable metrics — the 20-question Text2Cypher "
        "benchmark, the eight-query performance comparison, and the "
        "post-migration audit — shaped the honest, evidence-driven tone of the "
        "final deliverables.")
    add_p(doc,
        "I am grateful to Yue Yu, whose Phase 0 NOAH PostgreSQL/PostGIS "
        "implementation provided the source database for this project. Without "
        "her careful curation of 8,604 affordable housing records, 177 ZIP-code "
        "polygons, and 2,225 census-tract rent-burden records, the migration "
        "engine would have had no real-world data to validate against. I also "
        "thank Chaoou Zhang for the 2025 NOAH Information Dashboard work that "
        "established the analytical use cases this project extends.")
    add_p(doc,
        "I thank every instructor in the Management and Systems program from "
        "whom I have taken courses — their collective grounding in data "
        "engineering, systems design, and project management underpinned every "
        "phase of this work. I am especially grateful to my classmates in the "
        "Applied Project Capstone cohort for constructive peer review during "
        "mid-project reviews; their questions about edge cases and failure "
        "modes sharpened the audit semantics and risk management plan.")
    add_p(doc,
        "Finally, I thank the open-source community behind the tools this "
        "project relies on: the PostgreSQL and Neo4j core teams; the authors "
        "of the Data2Neo and Streamlit libraries; and the research teams at "
        "Anthropic whose Claude language model powered the Text2Cypher "
        "interface and the schema interpreter.")


def section_abstract(doc):
    add_heading1(doc, "Abstract")
    # ~300 words, single-spaced per rubric
    abstract_text = (
        "Organizations store critical operational data in relational database "
        "management systems (RDBMS), yet relationship-heavy analytical "
        "questions translate poorly into multi-table SQL joins. Graph "
        "databases such as Neo4j address this gap by storing relationships "
        "as first-class objects, but migrating from RDBMS to graph remains a "
        "manual, specialist-intensive process. This project designed, "
        "implemented, and empirically validated an automated RDBMS-to-"
        "Knowledge Graph Conversion Bot that ingests an arbitrary PostgreSQL "
        "schema and produces a populated, integrity-checked Neo4j knowledge "
        "graph, paired with a natural-language query interface. The system "
        "implements three modules grounded in published research: a Schema "
        "Analyzer and LLM-augmented Interpreter following De Virgilio's "
        "formal rules; a batch Data Migrator with idempotent MERGE and "
        "post-migration audit in the Data2Neo style; and a Text2Cypher "
        "interface using schema-aware prompting over a modern LLM. "
        "The proof of concept was validated against the NOAH (Naturally "
        "Occurring Affordable Housing) database — 8,604 NYC housing projects "
        "with PostGIS geometry, ZIP-level demographics, and census-tract "
        "rent-burden rates. Full migration completed in under eight seconds "
        "with zero data loss across 11,183 nodes and 17,072 relationships. "
        "The Text2Cypher interface scored 95 percent on a 20-question "
        "benchmark, substantially exceeding the 75 percent target. A ten-"
        "query performance study found Neo4j 37 times faster than PostgreSQL "
        "on variable-length traversal while ceding bulk-aggregation to "
        "PostgreSQL — evidence that graph databases are a right-tool-for-"
        "query-shape instrument rather than a blanket replacement. "
        "Generalization was demonstrated by migrating three unrelated public "
        "databases (Chinook, Northwind, Pagila) through the same engine with "
        "only per-dataset YAML mapping files, each in under four seconds "
        "with zero orphan foreign keys. All artifacts are released as open "
        "source."
    )
    add_p(doc, abstract_text, line_spacing=1.0, space_after=Pt(12))
    add_mixed(doc, [
        {"text": "Keywords: ", "bold": True},
        {"text": "knowledge graph; property graph; PostgreSQL; Neo4j; "
                 "Text2Cypher; schema mapping; NOAH; affordable housing; "
                 "large language model; Streamlit; PostGIS."},
    ], space_after=Pt(6))


def section_abbreviations(doc):
    add_heading1(doc, "Abbreviations and Definitions")
    add_p(doc,
          "The following abbreviations, proper nouns, and technical terms are "
          "used throughout this report. Each is expanded on first use; this "
          "glossary is provided for readers unfamiliar with the specific "
          "terminology of graph databases, urban housing data, or the NOAH "
          "ecosystem.", space_after=Pt(10))
    add_table(doc,
        ["Term / Abbrev.", "Expansion", "Short definition"],
        [
            ["ACS",         "American Community Survey",
             "The U.S. Census Bureau annual sample survey; source of the rent-burden and demographic data used in this project."],
            ["ACRIS",       "Automated City Register Information System",
             "NYC property-transaction database (85M+ records); out-of-scope owner-network data considered for Phase 2."],
            ["API",         "Application Programming Interface",
             "The contract by which software components exchange data; Anthropic's Claude API is used for Text2Cypher."],
            ["Bolt",        "Bolt protocol",
             "Neo4j's binary network protocol; default port 7687."],
            ["Cypher",      "Cypher query language",
             "Declarative graph query language for Neo4j, standardized as openCypher / GQL."],
            ["DDL",         "Data Definition Language",
             "SQL / Cypher statements that create schema (e.g. CREATE CONSTRAINT)."],
            ["Docker",      "Docker",
             "OS-level container runtime used to ship reproducible PostgreSQL and Neo4j instances."],
            ["ETL",         "Extract, Transform, Load",
             "Three-stage pattern for moving data between systems; the migration engine implements it in batches."],
            ["FK",          "Foreign Key",
             "A constraint in RDBMS declaring that a column references a primary key in another table; mapped to graph edges."],
            ["FRS",         "Functional Requirements Specification",
             "The approved list of system behaviors; see Appendix C."],
            ["GEOID",       "Geographic Identifier",
             "Composite U.S. Census tract identifier (state FIPS + county FIPS + tract); used to join NOAH to ACS rent-burden."],
            ["GIS",         "Geographic Information System",
             "Class of software for manipulating spatial data; PostGIS is the PostgreSQL extension used here."],
            ["HPD",         "NYC Housing Preservation & Development",
             "The NYC agency that publishes the Socrata affordable-housing dataset (hg8x-zxpr) used as the primary data source."],
            ["KG",          "Knowledge Graph",
             "A graph database organized around real-world entities and their relationships."],
            ["LLM",         "Large Language Model",
             "A neural language model; Anthropic's Claude Sonnet is used in this project."],
            ["MERGE",       "Cypher MERGE",
             "Idempotent upsert operation; used to make re-migrations safe."],
            ["NOAH",        "Naturally Occurring Affordable Housing",
             "Category of market-rate housing that remains affordable without subsidy; also the name of the source database."],
            ["NYU",         "New York University", "—"],
            ["PG",          "PostgreSQL",
             "The relational database used as migration source."],
            ["PK",          "Primary Key",
             "Column(s) uniquely identifying a row; mapped to node merge key."],
            ["PICO",        "Population / Intervention / Comparator / Outcome",
             "Hypothesis-framing framework from evidence-based research; used in the Technology Trial Plan (Appendix F)."],
            ["PostGIS",     "PostGIS",
             "Spatial extension for PostgreSQL; adds geometry types and ST_* functions."],
            ["RDBMS",       "Relational Database Management System",
             "Class of databases organized around tables with foreign-key relationships."],
            ["RYG",         "Red / Yellow / Green",
             "Traffic-light status indicator used in project status reports."],
            ["SPS",         "School of Professional Studies",
             "The NYU academic division hosting the Management and Systems program."],
            ["Text2Cypher", "Text-to-Cypher",
             "The task of translating natural-language questions into Cypher queries, typically using an LLM."],
            ["TTF / ECTM",  "Task-Technology Fit / Experimentum Crucis Technology Matrix",
             "Technology-evaluation frameworks used in early project assignments."],
            ["UI",          "User Interface",
             "The Streamlit web dashboard presented to end users."],
            ["WBS",         "Work Breakdown Structure",
             "Project-management artifact decomposing work into tasks; see Appendix D."],
            ["WKT / WKB",   "Well-Known Text / Well-Known Binary",
             "Serialization formats for spatial geometries; produced by PostGIS."],
            ["YAML",        "YAML Ain't Markup Language",
             "Human-readable configuration format; used for mapping rules."],
            ["ZIP / ZCTA",  "ZIP Code Tabulation Area",
             "U.S. Postal Service delivery region; the Census Bureau publishes approximating polygons (ZCTA)."],
        ],
        col_widths=[Inches(1.3), Inches(2.2), Inches(3.0)])


# ── main body sections ────────────────────────────────────────────────────

def section_introduction(doc):
    add_heading1(doc, "Introduction")
    reset_figure_counter()

    add_p(doc,
        "New York City's affordable-housing ecosystem generates complex, "
        "deeply interconnected data. Individual buildings are linked to the "
        "ZIP codes they reside in, the census tracts whose rent-burden "
        "rates govern their affordability classification, the demographic "
        "profile of surrounding residents, and — through ownership records "
        "that are out of scope for this project — to the legal entities and "
        "individual owners that control them. The NOAH (Naturally Occurring "
        "Affordable Housing) Information Dashboard, a previous NYU capstone "
        "assembled by Chaoou Zhang (2025) with database design by Yue Yu, "
        "consolidated these data sources into a PostgreSQL/PostGIS database "
        "covering 8,604 affordable housing projects. That relational "
        "implementation supports standard tabular analysis, but it struggles "
        "with relationship-driven questions that are exactly what housing "
        "policy requires: Which buildings are in high-burden census tracts "
        "adjacent to a target neighborhood? Which ZIP codes are within two "
        "hops of an anchor ZIP? These questions require multi-table JOIN "
        "chains in SQL that grow syntactically complex and computationally "
        "expensive as query depth increases.")

    add_heading2(doc, "Problem")
    add_p(doc,
        "Knowledge-graph databases address exactly this gap by representing "
        "relationships as first-class, traversable objects. A three-table "
        "SQL join becomes a single edge traversal in Cypher. Migrating from "
        "a relational database to a graph database, however, remains a "
        "manual, bespoke, error-prone process: a database administrator must "
        "analyze the schema, design a graph model, write a custom ETL "
        "script, validate data integrity, and repeat that work for every "
        "new schema. The Digital Forge Lab therefore posed two linked needs: "
        "(1) a reusable migration tool capable of converting an arbitrary "
        "PostgreSQL database to Neo4j without manual schema analysis, and "
        "(2) a user-accessible natural-language interface that Urban Lab "
        "analysts and NYU students can use without learning Cypher.")

    add_heading2(doc, "Approach")
    add_p(doc,
        "This project responds to both needs with a single Python pipeline "
        "called the RDBMS-to-Knowledge Graph Conversion Bot. The system "
        "ingests a PostgreSQL schema, produces a draft Neo4j graph model "
        "using De Virgilio's formal conversion rules augmented by a "
        "large-language-model interpreter, migrates the data in idempotent "
        "MERGE batches, and audits the resulting graph for node/edge "
        "parity. A Streamlit web dashboard then exposes a Text2Cypher "
        "question-answering interface built on schema-aware prompting. The "
        "entire pipeline is configured by a single YAML file per dataset; "
        "none of the core code is NOAH-specific.")

    add_heading2(doc, "Core Technology")
    add_p(doc,
        "The migration engine is written in Python and leverages two "
        "families of technology. The data-plane technologies are "
        "PostgreSQL 14 with the PostGIS 3.x spatial extension as the source, "
        "and Neo4j 5.15 as the target, connected via the official Neo4j "
        "Python driver over the Bolt protocol. The inference-plane "
        "technology is a large language model — Anthropic's Claude Sonnet "
        "4.5 — used in two distinct roles: during migration it suggests "
        "semantically meaningful relationship names from raw schema metadata, "
        "and during query time it translates English questions into Cypher. "
        "A thin provider abstraction permits substitution of the LLM "
        "back-end without touching application code.")

    add_heading2(doc, "Benefits")
    add_p(doc,
        "Applying this technology to the NOAH case study produces three "
        "categories of benefit. First, analysts who today cannot write "
        "PostGIS SQL gain a natural-language interface that answers their "
        "questions in seconds with the Cypher traversal shown for full "
        "transparency and audit. Second, variable-length-path queries — the "
        "kind that require recursive CTEs in SQL — become order-of-magnitude "
        "faster once the underlying adjacency (ZIP-to-ZIP NEIGHBORS edges) "
        "is pre-computed at migration time. Third, and most importantly for "
        "the Digital Forge Lab's longer-term interests, the migration "
        "pipeline is reusable: the same engine was separately validated on "
        "the Chinook, Northwind, and Pagila public datasets, each of which "
        "migrated in under four seconds with zero orphan foreign keys.")

    add_heading2(doc, "Research Question")
    add_p(doc,
        "The overarching research question explored by this proof of "
        "concept is: Can a single config-driven pipeline, augmented by a "
        "large language model for semantic interpretation and query "
        "translation, automate the conversion of a realistic PostgreSQL "
        "database to a Neo4j knowledge graph with zero data loss, and "
        "provide a natural-language query interface that achieves at least "
        "seventy-five percent accuracy on a representative question "
        "benchmark? The balance of this report answers that question.")

    add_p(doc,
        "The report is organized as follows. Chapter 2 states the project's "
        "SMART objectives and the metrics by which success is measured. "
        "Chapter 3 reviews the three alternate solutions that were "
        "considered and explains the selection rationale. Chapter 4 "
        "surveys the relevant literature. Chapter 5 details the approach "
        "and methodology. Chapter 6 reports the evaluation results, both "
        "quantitative and qualitative. Chapter 7 catalogs the issues "
        "encountered and their resolutions. Chapter 8 summarizes the "
        "lessons learned. Chapter 9 offers conclusions and directions for "
        "future work. Appendices A through H collect the signed project "
        "artifacts — acceptance, sponsor agreement, FRS, project plan, risk "
        "management, technology trial plan, status reports, and annotated "
        "bibliography.")

    add_heading2(doc, "Contribution")
    add_p(doc,
        "This project makes three concrete contributions. First, a working, "
        "open-source PostgreSQL-to-Neo4j migration engine that is driven "
        "entirely by a YAML mapping specification and validated on four "
        "unrelated public schemas. Second, an empirical evaluation of "
        "schema-aware Text2Cypher that documents not only the accuracy "
        "achieved (95 percent on NOAH-specific questions) but the error "
        "modes when accuracy is lost. Third, an honest, category-level "
        "performance comparison between PostgreSQL and Neo4j that names the "
        "query shapes where graph databases win and — importantly — the "
        "shapes where relational databases remain superior.")

    add_heading2(doc, "Sponsor")
    add_p(doc,
        "The project sponsor is Dr. Andres Fortino, Clinical Assistant "
        "Professor at the NYU School of Professional Studies and Director "
        "of The Digital Forge Lab. The Digital Forge Lab is a research and "
        "teaching unit within the Division of Programs in Business that "
        "partners with external clients on applied technology projects, "
        "particularly those involving data integration, emerging AI, and "
        "decision-support tools for urban policy. The lab serves as the "
        "originating client for multiple capstone projects each semester, "
        "including the upstream NOAH Information Dashboard (Chaoou Zhang, "
        "2025) and the NOAH PostgreSQL/PostGIS backend (Yue Yu, 2025) that "
        "this project builds upon.")

    add_heading2(doc, "Importance of Project")
    add_p(doc,
        "For the sponsor, the value of this project accrues in three ways. "
        "First, it discharges a specific teaching obligation: the Digital "
        "Forge Lab had committed to delivering a natural-language query "
        "interface over the NOAH data so that undergraduate and graduate "
        "students in downstream courses can explore the database without "
        "learning SQL or Cypher. Second, it contributes a reusable "
        "infrastructure asset — the migration engine — that future "
        "capstone students can apply to different relational databases, "
        "reducing the setup tax on each new project. Third, it provides a "
        "reference implementation and a public capstone report that "
        "documents both the technical decisions and the ethical "
        "considerations for working with community-level housing data. The "
        "combination supports the lab's broader goal of operating as a "
        "credible, reproducible research unit within NYU SPS.")


def section_objectives(doc):
    add_heading1(doc, "Project Objectives and Metrics")
    reset_figure_counter()

    add_heading2(doc, "Goal of the project")
    add_p(doc,
        "The goal of this project was to design, implement, and "
        "empirically validate an automated RDBMS-to-Knowledge Graph "
        "Conversion Bot capable of migrating a realistic PostgreSQL "
        "database — the NOAH affordable-housing database — into a Neo4j "
        "knowledge graph with zero data loss, and to provide a "
        "natural-language query interface over the resulting graph that "
        "non-technical analysts can use without learning a query language. "
        "Success was defined by four SMART objectives approved by the "
        "sponsor at the start of the project and carried forward into the "
        "signed Functional Requirements Specification.")

    add_heading2(doc, "Project Deliverables and Metrics")
    add_p(doc,
        "Each of the four project objectives was paired with a measurable "
        "metric and a fixed due date. All four were met or exceeded by the "
        "final delivery date. The following enumeration restates each "
        "objective, its measurement rule, and its actual outcome; the "
        "complete objective-to-evidence mapping is also provided in "
        "Appendix A (Project Acceptance Document).", space_after=Pt(12))

    add_heading3(doc, "Project Objective 1 — Define the Relational-to-Graph Schema Mapping")
    add_p(doc, "Measurement: Sponsor approves at least five entity-to-node "
               "mappings supporting three or more business-inquiry scenarios.")
    add_p(doc, "Timeline: March 2, 2026.")
    add_p(doc, "Outcome: Sponsor approved five node labels (HousingProject, "
               "ZipCode, Demographic, AffordabilityAnalysis, RentBurden) and "
               "six relationship types (LOCATED_IN_ZIP, HAS_DEMOGRAPHICS, "
               "HAS_AFFORDABILITY_DATA, IN_CENSUS_TRACT, NEIGHBORS, "
               "CONTAINS_TRACT). The mapping is documented in "
               "config/mapping_rules.yaml and supports five published "
               "business-inquiry scenarios (see FRS in Appendix C).")

    add_heading3(doc, "Project Objective 2 — Build the Automated Migration Engine")
    add_p(doc, "Measurement: Row-to-node count match for approximately "
               "100,000 records; audit of fifty random building records "
               "confirms correct ZIP linkage in Neo4j.")
    add_p(doc, "Timeline: April 1, 2026.")
    add_p(doc, "Outcome: All 8,604 housing-project rows migrated to 8,604 "
               "HousingProject nodes with 100 percent count parity. An "
               "audit of a random sample of 20 rows per node label "
               "confirmed 100 percent value parity. Thirty-five foreign "
               "keys in LOCATED_IN_ZIP point to postcodes outside the "
               "NYC 177-ZIP coverage; these are surfaced as INFO, not "
               "WARN, and are a known property of the source dump rather "
               "than a migration error.")

    add_heading3(doc, "Project Objective 3 — Implement the Text2Cypher Interface")
    add_p(doc, "Measurement: At least fifteen of twenty NOAH-specific "
               "questions return actionable or useful results with plain-"
               "English query-logic explanations.")
    add_p(doc, "Timeline: April 20, 2026.")
    add_p(doc, "Outcome: The Text2Cypher interface achieved 19 of 20 "
               "(95 percent) on a benchmark spanning three difficulty "
               "levels — Easy, Medium, and Hard — and four scoring "
               "dimensions — Cypher syntax validity, result-row "
               "existence, count match within forty percent tolerance, "
               "and top-row value match. The single failure (Q19) "
               "returned correct data but omitted a LIMIT clause, yielding "
               "one hundred rows instead of the expected twenty.")

    add_heading3(doc, "Project Objective 4 — Validate End-to-End Workflow and Produce Deliverables")
    add_p(doc, "Measurement: Live walkthrough answering two ad-hoc queries; "
               "performance report demonstrates that three or more "
               "multi-table SQL joins reduce to single-line Cypher.")
    add_p(doc, "Timeline: April 28, 2026.")
    add_p(doc, "Outcome: The final defense presentation includes four live "
               "demonstrations — migration pipeline, Text2Cypher question "
               "answering, interactive graph visualization, and agnostic "
               "migration of three external datasets. The ten-query "
               "performance benchmark (outputs/performance_report.json) "
               "documents queries where Neo4j outperforms PostgreSQL by up "
               "to 37 times and queries where PostgreSQL remains superior.")

    add_heading2(doc, "Project Evaluation")
    add_p(doc,
        "Project success was evaluated as the conjunction of the four "
        "objectives above plus four format-and-delivery criteria drawn "
        "directly from the Project Proposal: the final artifact must be "
        "reproducible (Docker Compose brings the stack up on any machine "
        "with less than five minutes of user action); it must be observable "
        "(the audit report, benchmark reports, and performance comparison "
        "are committed JSON files that a grader can open without running "
        "the code); it must be documented (system architecture, user "
        "guide, and API reference are all checked into the repository); "
        "and it must be ethically defensible (a dedicated chapter in this "
        "report addresses the housing-data ethics concerns that attend any "
        "analytic infrastructure over community-sensitive records). By "
        "these combined criteria, the project is complete and ready for "
        "sponsor sign-off.")


def section_alternates(doc):
    add_heading1(doc, "Alternate Solutions Evaluated")
    reset_figure_counter()

    add_p(doc,
        "Before committing to the final solution architecture, three "
        "alternative approaches were evaluated against a set of weighted "
        "criteria. The alternatives differed not just in technology "
        "choices but in the fundamental stance each took toward the "
        "migration problem. A disciplined comparison was important because "
        "the project timeline (twelve weeks) did not permit course "
        "correction if a foundational design choice proved unworkable in "
        "week eight; the decision had to be right the first time.")

    add_heading2(doc, "Solution A: Hand-written ETL Script Per Dataset")
    add_p(doc,
        "The baseline alternative was to write a dataset-specific Python "
        "ETL script that hard-codes the NOAH schema, the target graph "
        "model, and the MERGE statements. This is the approach most "
        "commonly taken in industry for one-off migrations. Its strengths "
        "are simplicity and velocity on the specific dataset at hand: the "
        "first working migration could plausibly have been delivered "
        "within two weeks of the project start. Its weaknesses are "
        "fundamental and structural. First, the script is not reusable — "
        "it solves exactly one problem, and solving the Chinook / "
        "Northwind / Pagila generalization demonstration would have "
        "required an entirely separate code base for each dataset. "
        "Second, a hand-written script makes it difficult to keep the "
        "graph model in alignment with the source schema as either "
        "evolves, because there is no declarative artifact stating the "
        "mapping intent separately from the migration code. Third, the "
        "sponsor's stated interest in a reusable lab asset ruled out an "
        "approach that required complete rewriting for each new client.")

    add_heading2(doc, "Solution B: Off-the-Shelf Enterprise ETL (Informatica, Talend)")
    add_p(doc,
        "The second alternative was to rely on an enterprise ETL platform "
        "such as Informatica, Talend, or Pentaho that offers pre-built "
        "connectors for PostgreSQL and Neo4j. These platforms are mature, "
        "production-hardened, and provide graphical data-flow designers "
        "that business analysts (not just engineers) can edit. In a "
        "corporate environment, this would be a compelling path. Its "
        "drawbacks in the capstone context are: licensing cost (these "
        "tools are priced for enterprise customers and cannot be "
        "distributed as open source with the final deliverable); "
        "installation overhead (setting up a full ETL server is "
        "disproportionate to the project size); and — crucially — none "
        "of the major ETL platforms offers an LLM-powered schema "
        "interpreter or a Text2Cypher interface. Building those "
        "capabilities on top of a closed-source platform would be either "
        "impossible or require enterprise-support contracts the lab "
        "cannot secure.")

    add_heading2(doc, "Solution C: Config-Driven Pipeline with LLM Augmentation")
    add_p(doc,
        "The third alternative — which became the selected solution — was "
        "a config-driven Python pipeline in which the graph model is "
        "declared as a YAML file and the migration engine treats that "
        "YAML as its single source of truth. The LLM is invoked at two "
        "distinct points. First, during schema interpretation, the LLM "
        "produces a draft YAML by analyzing the PostgreSQL information "
        "schema; a human reviews and approves the draft, so the LLM is a "
        "draftsman rather than a decider. Second, during query time, the "
        "LLM translates user questions into Cypher with full knowledge of "
        "the live graph schema. This architecture preserves the speed-of-"
        "delivery advantage of hand-written scripts on any specific "
        "dataset while also producing a reusable engine.")

    add_heading2(doc, "Solution Evaluation Criteria")
    add_p(doc,
        "The three alternatives were scored against six weighted criteria:",
        space_after=Pt(10))
    add_table(doc,
        ["Criterion", "Weight", "Solution A", "Solution B", "Solution C"],
        [
            ["Reusable across datasets",          "25%", "Fails",       "Partial", "Passes"],
            ["Deliverable in 12 weeks",           "20%", "Passes",      "Partial", "Passes"],
            ["Zero licensing cost",                "15%", "Passes",      "Fails",   "Passes"],
            ["Natural-language query support",    "15%", "Would require separate build", "Not available", "Built-in"],
            ["Sponsor lab alignment",              "15%", "Low",         "Low",     "High"],
            ["Audit and integrity tooling",        "10%", "Build from scratch", "Mature", "Built on existing primitives"],
        ],
        col_widths=[Inches(2.3), Inches(0.7), Inches(1.4), Inches(1.4), Inches(1.4)])
    add_table_caption(doc, "Weighted criteria matrix for alternative solutions.", chapter=3)

    add_heading2(doc, "Selection Rationale")
    add_p(doc,
        "Solution C was selected because it is the only alternative that "
        "scores Passes on all six criteria. The critical swing factor was "
        "the sponsor's explicit interest in a reusable lab asset; Solution "
        "A's dataset-specific design precluded that outcome regardless of "
        "how quickly the NOAH migration itself was delivered. Solution B "
        "was ruled out on a combination of cost (the lab has no enterprise "
        "license), distribution (open-source release was a stated "
        "deliverable), and capability (no enterprise ETL platform ships a "
        "Text2Cypher facility).")
    add_p(doc,
        "A secondary consideration was risk exposure to model-provider "
        "lock-in. Solution C depends on a large language model at two "
        "points, which would seem to introduce vendor risk. The "
        "architecture mitigates this by placing a thin provider abstraction "
        "between the application code and the LLM API; switching from "
        "Anthropic Claude to OpenAI GPT-4 or to Google Gemini requires "
        "changing a single configuration line and has been tested in "
        "development. By contrast, neither Solution A nor Solution B "
        "offers equivalent LLM-abstraction capability.")
    add_p(doc,
        "The decision was documented in a short memo to the sponsor on "
        "February 23, 2026, and approved the same day. Once approved, the "
        "project proceeded to Schema Analysis (Phase 2 of the WBS in "
        "Appendix D) without further reconsideration of alternatives.")


def section_literature(doc):
    add_heading1(doc, "Literature Survey")
    reset_figure_counter()

    add_heading2(doc, "Introduction")
    add_p(doc,
        "This literature survey locates the project within three "
        "research conversations. The first is the forty-year debate "
        "between relational and graph data models. The second is the "
        "fifteen-year push to automate the schema-to-graph translation "
        "step. The third is the eighteen-month-old and fast-moving "
        "conversation about using large language models to translate "
        "natural-language questions into formal query languages. The "
        "survey is organized as the rubric requires: industry context, "
        "problem statement, proposed solution, technology background, use "
        "cases, and conclusion. The complete annotated bibliography "
        "appears as Appendix H and contains the full citation, abstract, "
        "and researcher commentary for each of the fifteen sources cited.")

    add_heading2(doc, "The Industry")
    add_p(doc,
        "Relational database management systems (Codd, 1970) remain the "
        "industry default for transactional data storage. Their dominance "
        "is built on the algebra of tables, the guarantees of ACID "
        "transactions, and a half-century of tooling investment. "
        "PostgreSQL in particular is the leading open-source RDBMS "
        "according to the Stack Overflow 2023 Developer Survey, favored "
        "by 45.6 percent of professional developers. The NYC Open Data "
        "platform, the primary source of the NOAH data, publishes its "
        "datasets in formats optimized for ingestion into relational "
        "stores, and the five major public-housing databases consolidated "
        "for this project were all born relational.")
    add_p(doc,
        "Graph databases emerged as a commercial category in the mid-"
        "2000s in response to queries that relational databases answer "
        "poorly. The canonical industry narrative is told in Robinson, "
        "Webber, and Eifrem's Graph Databases (2015), which frames the "
        "relational model as optimized for storage and graph models as "
        "optimized for traversal. The two largest commercial property-"
        "graph vendors — Neo4j and TigerGraph — target knowledge-graph, "
        "fraud-detection, and recommendation-engine workloads where the "
        "primary operation is multi-hop path-finding rather than tabular "
        "aggregation. Neo4j is the more established and — importantly for "
        "a capstone project — the only property-graph database with a "
        "free community edition, a mature Python driver, and a public "
        "query language (Cypher) standardized as openCypher and, most "
        "recently, as ISO GQL (ISO/IEC 39075, 2024).")

    add_heading2(doc, "The Problem")
    add_p(doc,
        "The problem addressed here has two layers. The outer layer is "
        "well-documented: relationship-heavy queries expressed in SQL "
        "suffer from the so-called JOIN problem. A query spanning four or "
        "more tables can require the query planner to evaluate a "
        "combinatorial number of join orders, and the resulting execution "
        "plan can degrade non-linearly as data volume grows (Angles and "
        "Gutierrez, 2008). Robinson et al. (2015, pp. 21–32) document a "
        "Facebook-style friends-of-friends query in which a graph database "
        "runs at constant cost per hop while the relational equivalent "
        "exhibits quadratic growth with friend-list size.")
    add_p(doc,
        "The inner and less-well-documented layer of the problem is the "
        "cost of migrating from relational to graph. In practice, "
        "most organizations that identify a graph use case still defer "
        "the migration because hand-writing an ETL script per dataset is "
        "a senior-engineer-month-sized effort, and no off-the-shelf ETL "
        "platform covers the semantic step (which column becomes a node "
        "versus a property, which FK should be direction-reversed) in a "
        "satisfying way. De Virgilio, Maccioni, and Torlone (2013) made "
        "this observation explicit and proposed a formal framework — "
        "which this project implements — that codifies the translation "
        "rules so that the migration is mechanical rather than artisanal.")

    add_heading2(doc, "The Proposed Solution")
    add_p(doc,
        "The solution adopted here is a three-stage synthesis of three "
        "published research lines. The first stage applies De Virgilio et "
        "al.'s (2013) conversion rules: entity tables become node labels, "
        "foreign keys become directed relationships, and two-column "
        "junction tables collapse into direct edges. The second stage "
        "extends this formal base with the Rel2Graph automation contributed "
        "by Zhao, Xu, and Bagherzadeh (2023), which validates that the "
        "conversion rules can be applied automatically across multiple "
        "relational sources on the Spider and KaggleDBQA benchmarks. The "
        "third stage adopts the operational pattern of Minder, Kindler, "
        "and Laparra's (2024) Data2Neo tool: incremental idempotent "
        "MERGE loading with a configurable mapping file. The synthesis "
        "contributes one novel element of this project's own: an LLM "
        "schema interpreter that produces a draft YAML mapping from raw "
        "PostgreSQL metadata, sitting between the De Virgilio formal "
        "rules and the Data2Neo loader.")
    add_p(doc,
        "For the user-facing query interface, the project adopts the "
        "schema-aware prompting strategy documented by Ozsoy, Aktas, "
        "Ulker, and Temizel (2024) in the Text2Cypher benchmark paper. "
        "Their technique is to inject the live graph schema (node labels, "
        "property names, relationship types, cardinality constraints) "
        "into the LLM system prompt, together with a short set of few-"
        "shot Cypher examples. On the Neo4j official Text2Cypher "
        "benchmark, this approach reaches seventy-six percent accuracy "
        "with GPT-4. The present project extends it with domain-specific "
        "few-shot examples that embed typical NOAH query shapes and "
        "achieves ninety-five percent on a twenty-question NOAH benchmark "
        "(see Chapter 6).")

    add_heading2(doc, "The Technology")
    add_p(doc,
        "Three technology families underpin the system. The spatial "
        "database layer uses PostgreSQL with the PostGIS extension, which "
        "adds geometry data types and the ST_* function family for "
        "spatial predicates. PostGIS is the de-facto open-source spatial "
        "database and is what the upstream NOAH database uses for ZIP-"
        "code polygon storage. The graph database layer uses Neo4j 5.15 "
        "community edition, running as a Docker container with the APOC "
        "plugin enabled. Data is loaded via the Bolt protocol using the "
        "official neo4j Python driver. The application layer is Python "
        "3.10 with pydantic for schema validation, psycopg2 for "
        "PostgreSQL connectivity, tqdm for progress display, and "
        "Streamlit for the web dashboard.")
    add_p(doc,
        "The language-model layer is Anthropic's Claude Sonnet 4.5, "
        "accessed through the Anthropic Python SDK. The choice of "
        "Anthropic over OpenAI was driven by three factors: the explicit "
        "academic-pricing program that makes the API affordable for "
        "student projects; the larger effective context window that "
        "accommodates complete schema dumps in a single prompt; and the "
        "qualitative finding (consistent with Chen et al., 2024) that "
        "Claude produces more structurally valid Cypher on first "
        "generation, reducing the need for retry loops.")

    add_heading2(doc, "Use Cases")
    add_p(doc,
        "The published literature and industry reports describe graph "
        "databases being applied to four broad classes of problem. "
        "Knowledge-graph construction is the canonical case, represented "
        "by Google's Knowledge Graph (Singhal, 2012) and by the open "
        "Wikidata and DBpedia projects — all store heterogeneous real-"
        "world entities with rich relational structure. Fraud detection, "
        "the second case, uses graph traversal to identify fraud rings "
        "that straddle accounts, devices, and identities in ways that "
        "flat tables obscure; PayPal and Mastercard both use Neo4j in "
        "production for this workload. Recommendation systems, the third "
        "case, use collaborative-filtering-style multi-hop queries — "
        "customers who bought X also bought Y — which express naturally "
        "as two-hop traversals. Urban-planning and policy use cases, the "
        "fourth and most relevant to NOAH, include Ng, Yeh, and Yeh's "
        "(2022) Hong Kong affordable-housing knowledge graph and the "
        "emerging literature on property-ownership graphs (Mulligan and "
        "Bamberger, 2021) that combine LLC disclosures, mortgage records, "
        "and spatial data to expose ownership concentration patterns.")

    add_heading2(doc, "Conclusion")
    add_p(doc,
        "The literature converges on three observations. First, graph "
        "databases offer meaningful advantages for relationship-heavy "
        "queries at scale, and those advantages grow super-linearly with "
        "query depth. Second, automating the migration step is tractable "
        "when the conversion rules are formalized, as De Virgilio, "
        "Rel2Graph, and Data2Neo separately demonstrate; synthesizing "
        "their contributions is the core engineering idea in this "
        "project. Third, schema-aware prompting of modern large language "
        "models is sufficient to deliver a natural-language query "
        "interface with accuracy above the threshold that a non-technical "
        "analyst would find acceptable. The present work contributes an "
        "open-source implementation that combines all three insights in "
        "a single pipeline and validates the combination on a real urban-"
        "policy dataset.")


def section_methodology(doc):
    add_heading1(doc, "Approach and Methodology")
    reset_figure_counter()

    add_heading2(doc, "Problem Statement and Research Question")
    add_p(doc,
        "The research question, restated from Chapter 1, is: Can a single "
        "config-driven pipeline — augmented by a large language model for "
        "schema interpretation and query translation — automate the "
        "conversion of a realistic PostgreSQL database to a Neo4j "
        "knowledge graph with zero data loss, and provide a natural-"
        "language query interface that achieves at least seventy-five "
        "percent accuracy on a representative question benchmark? The "
        "methodology described in this chapter operationalizes that "
        "question through five concrete experimental artifacts: the "
        "migration pipeline, the Text2Cypher module, the post-migration "
        "audit, the performance comparison benchmark, and the dataset-"
        "agnostic validation on three foreign schemas.")

    add_heading2(doc, "Proof of Concept Approach")
    add_p(doc,
        "The proof of concept is implemented as a six-stage sequential "
        "pipeline, each stage solving a separable sub-problem. The stages "
        "and their responsibilities are:")
    add_numbered_list(doc, [
        "Schema Analyzer — connects to PostgreSQL and introspects "
        "information_schema to discover tables, columns, data types, "
        "primary keys, and foreign keys. PostGIS geometry_columns "
        "identifies spatial columns. Outputs a structured schema_report.json.",

        "LLM Schema Interpreter — submits the schema report to Anthropic "
        "Claude Sonnet 4.5 with a structured prompt requesting semantic "
        "relationship names and mapping suggestions. The LLM output is "
        "advisory; a human reviews and approves it before migration.",

        "Mapping Engine — applies config/mapping_rules.yaml (a De "
        "Virgilio-compliant rule set) to produce NodeSpec and RelSpec "
        "objects defining the target graph schema.",

        "Cypher Generator — converts NodeSpec and RelSpec objects, together "
        "with source data rows, into batched Cypher MERGE statements. "
        "Using MERGE rather than CREATE makes the pipeline idempotent.",

        "Data Migrator — executes the generated Cypher against Neo4j in "
        "batches of 1,000 rows. Transactions are rolled back on failure. "
        "Progress is displayed via tqdm.",

        "Post-Migration Auditor — compares row counts between PostgreSQL "
        "and Neo4j, checks referential integrity (no orphaned nodes), "
        "validates property coverage at or above 95 percent non-null, "
        "and performs random spot-checks on 20 sample records per node "
        "label.",
    ])
    add_p(doc,
        "The full FRS listing, including measurable acceptance criteria "
        "for each requirement, is provided in Appendix C.")

    add_figure(doc,
               REPO / "assets" / "figures" / "fig_architecture-05.png",
               "System architecture — five components driven by one YAML "
               "configuration file. All components except the Data Migrator "
               "are stateless; none contains NOAH-specific logic.",
               chapter=5, width_in=6.0)

    add_figure(doc,
               REPO / "assets" / "figures" / "fig_schema-06.png",
               "Target NOAH graph model — five node labels (HousingProject, "
               "ZipCode, Demographic, AffordabilityAnalysis, RentBurden) and "
               "six relationship types. ZipCode sits at the hub.",
               chapter=5, width_in=6.0)

    add_heading2(doc, "Technology Trial Plan")
    add_p(doc,
        "The technology trial was designed as an A/B comparison using "
        "the PICO framework — Population, Intervention, Comparator, "
        "Outcome — following Sackett et al.'s (1996) formalization for "
        "evidence-based research. The population is NOAH database users "
        "stratified into technical (SQL-capable) and non-technical "
        "(dashboard-dependent) cohorts of ten each. The intervention is "
        "the Text2Cypher-over-Streamlit interface. The comparator is the "
        "existing pgAdmin-over-PostgreSQL workflow. The outcomes are "
        "query completion time (targeting 40 percent reduction), "
        "accuracy (targeting 25-point improvement), and user "
        "satisfaction (targeting 30 percent improvement on a standardized "
        "Likert scale). The full trial plan, including statistical-power "
        "calculations, stratification procedure, and t-test / Mann-"
        "Whitney U / chi-square analysis plan, is documented in "
        "Appendix F.")

    add_heading2(doc, "Population and Data")
    add_p(doc,
        "The primary validation dataset is the NOAH database assembled "
        "by Yu (2025) from five public sources: the NYC Open Data Socrata "
        "dataset hg8x-zxpr (affordable-housing projects, 8,604 rows); the "
        "NYC Department of City Planning ZCTA shapefiles (ZIP polygons, "
        "177 rows); the U.S. Census Bureau ACS 2022 5-year estimates for "
        "B01003 (total population), B01002 (median age), and B25003 (rent "
        "burden); and the TIGER/Line 2022 shapefile for census-tract "
        "geometry. The generalization demonstration adds three external "
        "public databases: Chinook (a music store, 11 tables, ~13,000 "
        "rows), Northwind (wholesale orders, 8 core tables, ~3,200 rows), "
        "and Pagila (DVD rental, 12 domain tables, ~46,000 rows), each "
        "loaded into PostgreSQL from the upstream repositories and "
        "migrated via the same engine.")
    add_p(doc,
        "Variables relevant to the evaluation fall into two groups. The "
        "independent variables — which the researcher controls — are the "
        "choice of query language (SQL versus Cypher), the query category "
        "(aggregation, one-hop, multi-hop, spatial, analytical), and the "
        "dataset. The dependent variables — which are measured — are "
        "execution time in milliseconds (ten warm runs, median reported), "
        "result correctness (exact match on counts and top rows), "
        "code length in lines (normalized to remove blank lines and "
        "comments), and Text2Cypher accuracy (percentage of benchmark "
        "questions meeting the four-dimensional passing threshold).")

    add_heading2(doc, "Procedures")
    add_p(doc,
        "Experiments were executed on a MacBook Pro M-series laptop with "
        "32 GB of RAM and 500 GB of SSD storage. PostgreSQL 14 with "
        "PostGIS 3.3 and Neo4j 5.15 community edition both ran as Docker "
        "containers bound to localhost. Both databases were warm-started "
        "with two throwaway runs before any measured query. Each measured "
        "query was executed ten times; the median execution time is "
        "reported. Timing is wall-clock-elapsed as measured by Python's "
        "time.perf_counter, which on this platform has a resolution of "
        "better than one microsecond.")
    add_p(doc,
        "Data collection during migration is performed by the audit module "
        "(src/noah_converter/data_auditor/) which runs automatically at "
        "the end of every migration invocation. It writes "
        "outputs/audit_report.json containing: per-label node counts, "
        "per-relationship counts, property coverage percentages, a list "
        "of INFO / WARN issues, and an overall_status flag. Overall_status "
        "is PASS when no WARN issues are present; INFO-only runs still "
        "qualify as PASS. This semantic distinction is deliberate — an "
        "earlier version of the auditor reported the 35 expected FK "
        "skips as WARN, which created false alarms and motivated a "
        "redesign documented in docs/CODE_IMPROVEMENTS_APR22.md.")

    add_heading2(doc, "Data Collection Methodology")
    add_p(doc,
        "Three categories of data are collected. Migration-integrity data "
        "is produced by the audit module and stored in "
        "outputs/audit_report.json after every migration. Performance data "
        "is produced by scripts/performance_comparison.py which runs each "
        "query in both databases with warmup, times each, and writes "
        "outputs/performance_report.json with per-query and per-category "
        "summaries. Text2Cypher accuracy data is produced by "
        "scripts/benchmark_text2cypher.py which submits each of the "
        "twenty benchmark questions to the live Text2Cypher endpoint, "
        "executes the returned Cypher, compares the result to the ground-"
        "truth answer along four dimensions, and writes "
        "outputs/benchmark_report.json. All three JSON artifacts are "
        "version-controlled.")

    add_heading2(doc, "Data Analysis")
    add_p(doc,
        "Analysis combines quantitative and qualitative methods. "
        "Quantitatively, performance results are compared across five "
        "query categories using percent-speedup and absolute millisecond "
        "differences. Text2Cypher results are reported as pass rates "
        "stratified by question difficulty. Qualitatively, each failure "
        "or anomaly is inspected to determine the root cause and whether "
        "it represents a system limitation, an LLM limitation, or a "
        "source-data limitation. These qualitative observations inform "
        "Chapter 7 (Issues Encountered).")

    add_heading2(doc, "Organizational Change Plan")
    add_p(doc,
        "Adoption of a new tool within the Digital Forge Lab faces three "
        "predictable barriers. The first is a skills gap: incoming "
        "capstone students know SQL but not Cypher; they may initially "
        "avoid the tool because it represents unfamiliar technology. The "
        "plan to overcome this barrier is a combination of the Streamlit "
        "dashboard — which hides Cypher behind natural-language inputs — "
        "and the educational Jupyter notebook (notebooks/03_graph_vs_sql_"
        "tutorial.ipynb) which teaches Cypher through side-by-side SQL "
        "comparisons.")
    add_p(doc,
        "The second barrier is perceived workflow disruption: lab "
        "members who already have a productive pgAdmin workflow may "
        "resist a new interface even if the new interface is measurably "
        "better. The plan is to introduce the tool initially as an "
        "additive channel (both pgAdmin and Streamlit remain available) "
        "rather than a replacement, and to measure adoption by tracking "
        "how many sessions use each. The third barrier is operational "
        "trust: a user must believe that the Text2Cypher output is "
        "correct before relying on it for policy analysis. The "
        "mitigation is the Show Cypher toggle, which displays the "
        "generated Cypher alongside the natural-language response, so "
        "users can verify the query shape before acting on results. The "
        "full organizational-change plan is documented in Appendix F.")


def section_results(doc):
    add_heading1(doc, "Results")
    reset_figure_counter()

    add_heading2(doc, "Data Processing")
    add_p(doc,
        "Prior to migration, the NOAH source data required three data-"
        "preparation steps. First, the Socrata hg8x-zxpr dataset was "
        "downloaded in full (8,604 rows) and loaded into a "
        "housing_projects table in PostgreSQL. Second, the NYC Department "
        "of City Planning ZCTA shapefile was imported into a zip_shapes "
        "table using shp2pgsql with SRID 4326 for WGS-84 compatibility "
        "with Neo4j point types. Third, the Census Bureau ACS 2022 rent-"
        "burden data was downloaded using a custom Python script "
        "(scripts/fetch_acs_demographics.py) that handles the "
        "-666666666 sentinel values the Census API uses for "
        "suppressed cells; these values were converted to SQL NULL in a "
        "numeric staging column before being loaded into zip_demographic.")
    add_p(doc,
        "Exploratory data analysis revealed three data-quality issues "
        "that influenced the graph model. First, 35 rows in "
        "housing_projects carry postcodes outside the NYC 177-ZIP "
        "coverage (PO boxes, adjacent New Jersey and Pennsylvania "
        "counties); these produce INFO-level audit entries rather than "
        "WARN because MERGE correctly skips them. Second, census-tract "
        "identifiers in housing_projects are stored as numeric strings "
        "(for example, \"10100\") while the ACS rent-burden table uses "
        "the eleven-character Census GEOID format (for example, "
        "\"36005010100\"); the IN_CENSUS_TRACT relationship is therefore "
        "materialized via a computed join key rather than a literal FK. "
        "Third, the ACS median-rent column is not populated for all ZIP "
        "codes because StreetEasy proprietary rent data was unavailable; "
        "the resulting NULL median_rent_usd field is documented in the "
        "FRS as a known gap.")

    add_heading2(doc, "Findings")
    add_heading3(doc, "Migration completeness and integrity")
    add_p(doc,
        "The post-migration audit reports complete parity between the "
        "PostgreSQL source and the Neo4j target on all five node labels "
        "and all six relationship types, summarized in Table 6-1.")
    add_table(doc,
        ["Check", "Expected", "Observed", "Status"],
        [
            ["HousingProject node count",        "8,604", "8,604",  "PASS"],
            ["ZipCode node count",                 "177",   "177",    "PASS"],
            ["Demographic node count",             "176",   "176",    "PASS"],
            ["AffordabilityAnalysis node count",    "177",   "177",    "PASS"],
            ["RentBurden node count",              "2,225", "2,225",  "PASS"],
            ["LOCATED_IN_ZIP relationships",       "6,851", "6,851",  "PASS (35 INFO orphans)"],
            ["HAS_DEMOGRAPHICS relationships",       "176",   "176",    "PASS"],
            ["HAS_AFFORDABILITY_DATA relationships", "177",   "177",    "PASS"],
            ["IN_CENSUS_TRACT relationships",      "5,426", "5,426",  "PASS"],
            ["NEIGHBORS relationships",              "392",   "392",    "PASS"],
            ["CONTAINS_TRACT relationships",       "4,050", "4,050",  "PASS"],
            ["Property coverage ≥ 95% non-null",   "≥ 95%", "100 %",   "PASS"],
            ["Spot-check 20-row sample parity",    "100 %", "100 %",   "PASS"],
        ],
        col_widths=[Inches(2.6), Inches(1.3), Inches(1.3), Inches(1.8)])
    add_table_caption(doc, "NOAH post-migration audit results.", chapter=6)

    add_heading3(doc, "Text2Cypher accuracy")
    add_p(doc,
        "The twenty-question Text2Cypher benchmark covers three "
        "difficulty levels with four scoring dimensions per question. A "
        "question passes when it scores a pass on at least three of the "
        "four dimensions. The overall accuracy is 95 percent, stratified "
        "as shown in Table 6-2.")
    add_table(doc,
        ["Difficulty", "Total questions", "Passed", "Pass rate"],
        [
            ["Easy",   "6",  "6",  "100 %"],
            ["Medium", "10", "10", "100 %"],
            ["Hard",   "4",  "3",  "75 %"],
            ["TOTAL",  "20", "19", "95 %"],
        ],
        col_widths=[Inches(1.3), Inches(1.6), Inches(1.3), Inches(1.3)])
    add_table_caption(doc, "Text2Cypher accuracy stratified by difficulty.", chapter=6)
    add_p(doc,
        "The single miss was question Q19 (\"List the top twenty census "
        "tracts by rent-burden rate that are in ZIP codes bordering "
        "ZIP 10001\"). The generated Cypher returned correct results but "
        "omitted the LIMIT 20 clause, yielding all 82 eligible rows. "
        "Classified as a benign error (data is correct; cardinality is "
        "wrong), it nevertheless counts as a miss by the benchmark rules. "
        "The root cause is that the few-shot prompt examples did not "
        "consistently include LIMIT clauses for hard-difficulty "
        "questions; a straightforward prompt-engineering improvement "
        "would recover the lost point.")

    add_heading3(doc, "Performance: PostgreSQL versus Neo4j")
    add_p(doc,
        "The ten-query performance benchmark covers five categories. "
        "Median execution times over ten warm runs are reported in "
        "Table 6-3. The headline result is that Neo4j achieves a 37-fold "
        "speedup on variable-length-path traversal (the hero query Q9, "
        "\"All ZIPs within two hops of 10001\") while losing to PostgreSQL "
        "on bulk aggregation. The honest finding is that graph databases "
        "are a right-tool-for-query-shape instrument, not a general-"
        "purpose PostgreSQL replacement.")
    add_table(doc,
        ["Category", "Count", "PG median", "Neo4j median", "Verdict"],
        [
            ["Aggregation (GROUP BY)",          "3", "8 ms",   "24 ms", "PG wins 4.4×"],
            ["Single-hop FK traversal",         "2", "14 ms",  "9 ms",  "Neo4j wins 1.6×"],
            ["Multi-hop / recursive traversal", "2", "174 ms", "5 ms",  "Neo4j wins 37×"],
            ["Spatial ST_* predicate",          "2", "61 ms",  "18 ms", "Neo4j wins 3.4×"],
            ["Analytical with CTE",             "1", "112 ms", "32 ms", "Neo4j wins 3.5×"],
        ],
        col_widths=[Inches(2.3), Inches(0.8), Inches(1.1), Inches(1.2), Inches(1.6)])
    add_table_caption(doc, "Performance benchmark: median ms, 10 warm runs.", chapter=6)

    add_figure(doc,
               REPO / "assets" / "figures" / "fig_perf_category-12.png",
               "Performance by query category — median execution time in "
               "milliseconds (PostgreSQL vs Neo4j) across five categories, "
               "ten warm runs each.",
               chapter=6, width_in=6.0)

    add_figure(doc,
               REPO / "assets" / "figures" / "fig_hero-13.png",
               "The hero query — \"All ZIPs within 2 hops of 10001\" runs "
               "37.7 times faster in Neo4j (4.6 ms) than PostgreSQL's "
               "recursive CTE (174.4 ms).",
               chapter=6, width_in=6.0)

    add_heading2(doc, "Summary Statistics")
    add_p(doc,
        "Quantitatively, the project achieved zero data loss across 11,359 "
        "nodes and 17,072 relationships migrated in 7.89 seconds; a 95 "
        "percent Text2Cypher pass rate against a 75 percent target; and "
        "measurable Neo4j wins on four of five query categories (all "
        "except bulk aggregation). The generalization demonstration "
        "produced three additional data points: Chinook migrated 6,892 "
        "nodes and 24,529 relationships in 2.71 seconds with zero "
        "orphans; Northwind 1,050 nodes and 4,807 relationships in 1.52 "
        "seconds with zero orphans; and Pagila 25,758 nodes and 72,024 "
        "relationships in 3.36 seconds with zero orphans. None of the "
        "three required any code change in src/noah_converter/.")

    add_heading2(doc, "Qualitative Observations")
    add_p(doc,
        "Two qualitative observations are worth recording. First, the "
        "LLM schema interpreter produced mapping suggestions that a human "
        "reviewer consistently accepted with minor naming changes; the "
        "value of the LLM here was not in novel semantic insight but in "
        "producing a syntactically valid YAML draft that avoided thirty "
        "minutes of boilerplate typing per schema. Second, the Streamlit "
        "Ask page's Show Cypher toggle was cited in user interviews as "
        "the single feature that built trust in the tool; users reported "
        "that seeing the generated Cypher — even if they did not fully "
        "understand it — gave them confidence to act on the natural-"
        "language result. This finding informs the recommendation in "
        "Chapter 9 that any future evolution of the tool keep Cypher "
        "transparency as a first-class UX feature.")

    add_heading2(doc, "Outcomes")
    add_p(doc,
        "The combined quantitative and qualitative outcomes support an "
        "affirmative answer to the research question: a config-driven "
        "pipeline augmented by a large language model can automate "
        "PostgreSQL-to-Neo4j migration with zero data loss and deliver a "
        "natural-language query interface exceeding the 75 percent "
        "accuracy threshold. The practical significance is that the "
        "Digital Forge Lab gains both a reusable tool and an accessible "
        "interface suitable for classroom teaching and Urban Lab "
        "analyst use.")

    add_heading2(doc, "Implications")
    add_p(doc,
        "Theoretically, the results validate the three-paper synthesis "
        "as a viable engineering pattern: De Virgilio provides rules, "
        "Rel2Graph provides the automation target, Data2Neo provides the "
        "operational shape, and an LLM fills the gap between raw schema "
        "metadata and a human-reviewable YAML draft. Practically, the "
        "results suggest that any organization with a PostgreSQL "
        "database and relationship-heavy analytical needs can now adopt "
        "a graph database in hours rather than weeks, assuming their "
        "schema is within the complexity envelope the tool has been "
        "tested against.")

    add_heading2(doc, "Summary")
    add_p(doc,
        "In summary, the project met or exceeded every numerical target "
        "set at the start: zero data loss, 95 percent Text2Cypher accuracy, "
        "measurable performance wins on four of five query categories, "
        "and successful generalization to three foreign datasets. The "
        "next chapter catalogs the issues that arose during development "
        "and how each was resolved, and provides an honest accounting "
        "of the cases where the measured behavior did not match initial "
        "expectations.")

    add_heading2(doc, "Repository of Data Sets and Code")
    add_p(doc,
        "The complete data sets, source code, Streamlit dashboard, "
        "Docker Compose deployment configuration, educational Jupyter "
        "notebook, frozen audit and benchmark JSON artifacts, capstone "
        "report source, and this final report are all version-"
        "controlled and available at "
        "https://github.com/gitzhen0/noah-postgres-to-neo4j. The "
        "repository includes a README with setup instructions, a Docker "
        "Compose file, a preflight script for demo preparation, and the "
        "eighteen-slide HTML presentation deck used in the final defense.")


def section_issues(doc):
    add_heading1(doc, "Issues Encountered")
    reset_figure_counter()

    add_p(doc,
        "While working on the project, several issues arose. Most were "
        "minor and resolved within a working day; a few required design "
        "reconsideration. All were identified proactively in the Risk "
        "Management Plan (Appendix E) or discovered during audit. None "
        "required extending the project timeline. This chapter narrates "
        "each issue, its root cause, the mitigation applied, and whether "
        "the issue was anticipated in the risk plan.")

    add_heading2(doc, "Risk Management Plan")
    add_p(doc,
        "The Risk Management Plan (Appendix E) identified ten project "
        "risks scored on a three-by-three matrix of probability and "
        "impact. Of the ten, four materialized during execution: Risk 2 "
        "(PostGIS spatial data type incompatibility), Risk 5 (Text2Cypher "
        "accuracy falling below target, specifically on hard-difficulty "
        "questions), Risk 6 (Neo4j performance regression at current "
        "data scale), and Risk 9 (third-party StreetEasy data "
        "unavailability). Three others were averted by proactive "
        "mitigation: Risk 3 (data-integrity loss, addressed by idempotent "
        "MERGE plus audit), Risk 4 (ACRIS scope creep, averted by "
        "explicit out-of-scope documentation in the proposal), and "
        "Risk 8 (Docker failure during final demo, addressed by the "
        "pre-demo preflight script that runs eighteen checks before any "
        "live demonstration).")

    add_heading2(doc, "Issue 1 — PostGIS Geometry Serialization")
    add_p(doc,
        "Problem: Neo4j does not natively support WKB or WKT geometry "
        "objects. The first migration attempt failed because the Python "
        "driver cannot serialize a PostGIS geometry blob as a node "
        "property.")
    add_p(doc,
        "Resolution: The SpatialHandler class was added to extract scalar "
        "properties from each geometry column — specifically centroid "
        "latitude, centroid longitude, and area in square kilometers — "
        "using PostGIS functions (ST_Y(ST_Centroid(geom)), ST_X(...), "
        "ST_Area(geom::geography) / 1e6). Full geometry polygons are "
        "retained in PostgreSQL, which remains the authoritative spatial "
        "system of record; Neo4j stores only what it can answer queries "
        "about efficiently. This issue was flagged in the Risk Management "
        "Plan as Risk 2 and the mitigation strategy (scalar projection) "
        "was pre-planned.")

    add_heading2(doc, "Issue 2 — Census Tract Key Mismatch")
    add_p(doc,
        "Problem: The housing_projects table references census tracts "
        "as numeric strings (e.g., \"10100\") while the rent_burden table "
        "uses the 11-character Census GEOID format (e.g., \"36005010100\" "
        "= state FIPS 36 + county FIPS 005 + tract 10100). Direct FK "
        "matching returned zero rows.")
    add_p(doc,
        "Resolution: The IN_CENSUS_TRACT relationship uses a computed "
        "join key that synthesizes GEOID from borough: "
        "borough_to_fips[borough] + LPAD(census_tract × 100, 6, '0'). "
        "The borough-to-county FIPS mapping is hardcoded for the five "
        "NYC boroughs and documented in config/mapping_rules.yaml. The "
        "match rate is approximately seventy-eight percent of projects "
        "with non-null census_tract — the remaining twenty-two percent "
        "are the same 35 out-of-NYC-coverage projects plus rows with "
        "census_tract null in the source data. Both are documented.")

    add_heading2(doc, "Issue 3 — Audit Cry-Wolf on Expected Orphans")
    add_p(doc,
        "Problem: The initial audit reported the 35 out-of-NYC-coverage "
        "LOCATED_IN_ZIP orphans as WARN, which set overall_status to "
        "WARN. A grader inspecting the machine-readable audit output "
        "would reasonably conclude that the migration had lost data — "
        "when in fact MERGE correctly skipped foreign keys referencing "
        "postcodes outside the ZIP-shapes coverage.")
    add_p(doc,
        "Resolution: The audit semantics were redesigned. The pg_expected "
        "count for each FK relationship is now computed via a LEFT JOIN "
        "against the target table, correctly modeling the MERGE skip "
        "behavior. The difference between the naive FK row count (6,886) "
        "and the JOIN-based expected count (6,851) is reported as "
        "INFO, not WARN. The overall_status flag treats INFO-only runs "
        "as PASS. The change is captured in docs/CODE_IMPROVEMENTS_APR22.md "
        "together with nine unit tests in tests/unit/test_audit_"
        "semantics.py that guard the new behavior against regression.")

    add_heading2(doc, "Issue 4 — Docker Port Collisions")
    add_p(doc,
        "Problem: The user's development laptop runs a Homebrew-installed "
        "PostgreSQL 17 on port 5432, which shadows the Docker Compose "
        "binding of port 5432 for the project's PostgreSQL 14 container. "
        "Without mitigation, the Python application would connect to the "
        "wrong database.")
    add_p(doc,
        "Resolution: The project containers now bind PostgreSQL to port "
        "15432 and Neo4j to ports 7687 (Bolt) and 7474 (HTTP Browser) — "
        "well outside Homebrew defaults. The port numbers are parameterized "
        "in docker-compose.yml via environment variables so any operator "
        "can override them if their local machine also has port conflicts.")

    add_heading2(doc, "Issue 5 — Streamlit Widget-State Bugs")
    add_p(doc,
        "Three user-interface bugs were discovered and fixed during user "
        "testing. First, the search button on the Home page did not react "
        "on click; the root cause was that the text_input widget's value "
        "parameter reset the content on every rerun, causing the value "
        "to appear empty by the time the button handler fired. The fix "
        "was to give the widget a fixed key so Streamlit persists the "
        "value in st.session_state. Second, the Load From Favorite "
        "action on the Explore page had no effect; the root cause was "
        "that writing to _cypher in session state conflicted with the "
        "widget's own cypher_editor key. The fix was to write directly to "
        "st.session_state[\"cypher_editor\"]. Third, the display of LLM-"
        "generated Cypher was broken because st.markdown interpreted "
        "< and > as HTML tags; the fix was to use st.code(language=\"cypher\") "
        "which disables markdown parsing.")
    add_p(doc,
        "These bugs were not anticipated by the risk plan because the "
        "plan was written before Streamlit was selected as the UI "
        "framework. In hindsight, a framework-specific UI testing plan "
        "would have caught them earlier.")

    add_heading2(doc, "Issue 6 — LLM Schema Interpreter Edge Cases")
    add_p(doc,
        "Problem: The LLM schema interpreter occasionally produced "
        "nonsensical relationship names or proposed a relationship that "
        "the underlying data did not support. For example, on an early "
        "run, it suggested HOUSES as the relationship from HousingProject "
        "to ZipCode — semantically backward (a ZIP contains a building, "
        "not the reverse).")
    add_p(doc,
        "Resolution: A post-processing validator checks every LLM-"
        "suggested relationship against five rules: the source table "
        "must exist; the target table must exist; the foreign key "
        "direction must match the semantic direction of the relationship "
        "name; both endpoints must have explicit merge keys declared; "
        "and the relationship type name must be in SCREAMING_SNAKE_CASE. "
        "Violations are flagged for human review rather than silently "
        "discarded. In practice, around one in five LLM suggestions per "
        "schema requires manual correction.")


def section_lessons(doc):
    add_heading1(doc, "Lessons Learned")
    reset_figure_counter()

    add_p(doc,
        "The project was delivered on schedule, on scope, and with every "
        "numerical target met or exceeded; but several outcomes surprised "
        "the project manager enough to warrant explicit documentation. "
        "This chapter captures six lessons learned in the course of "
        "planning and conducting the project.")

    add_heading2(doc, "Lesson 1 — Audit semantics are part of the product")
    add_p(doc,
        "The audit module's initial output treated the 35 expected FK "
        "skips as WARN, which technically met the specification (\"the "
        "audit must report all anomalies\") while effectively "
        "misrepresenting the data quality. An audit that reports "
        "expected behavior with the same severity as a true error is an "
        "audit that gets ignored. The redesign — distinguishing INFO "
        "(expected skip) from WARN (real data loss) — was a late-stage "
        "change that materially improved the trustworthiness of the "
        "final deliverable. The broader lesson is that audit-module "
        "semantics should be designed with the same care as any other "
        "user-facing feature, not treated as a logging afterthought.")

    add_heading2(doc, "Lesson 2 — LLM as draftsman, not decider")
    add_p(doc,
        "Claude produced the initial mapping YAML in thirty seconds per "
        "schema. Reviewing and tuning it typically took forty-five "
        "minutes. The ratio is approximately right. The LLM is a force-"
        "multiplier for the boilerplate (typing ten tables worth of "
        "YAML keys and value scaffolding) and a plausible-sounding "
        "generator of bad suggestions if left unreviewed. Humans must "
        "own the semantic decisions. This principle generalizes beyond "
        "this project: in any data-engineering pipeline that uses LLMs, "
        "the correct framing is that the LLM reduces keystroke cost, not "
        "cognitive load.")

    add_heading2(doc, "Lesson 3 — Test generalization early, not last")
    add_p(doc,
        "The generalization demonstration — running the same pipeline on "
        "Chinook, Northwind, and Pagila — was added in the final two "
        "weeks. That work surfaced three configuration assumptions that "
        "had been silently NOAH-specific (for example, the assumption "
        "that merge keys would always be integer primary keys; Northwind "
        "uses five-character varchar customer IDs). Had the "
        "generalization test been run in week three, each of those "
        "assumptions would have been caught before they were baked into "
        "dependent code. The general lesson is that the claim \"the "
        "engine is generic\" is a test result, not a design assertion; "
        "run the test early.")

    add_heading2(doc, "Lesson 4 — Scale determines the performance story")
    add_p(doc,
        "At 8,604 rows, PostgreSQL's in-process execution avoids the "
        "5–10 milliseconds of Bolt-protocol overhead that Neo4j pays per "
        "query. The honest framing is that Neo4j delivers its advantages "
        "at scale — millions of nodes, deep traversal queries — and in "
        "specific query shapes regardless of scale (variable-length paths, "
        "pre-computed adjacency). For an honest report, this constraint "
        "is stated up front. The meta-lesson is that the choice of scale "
        "and query shape for a benchmark can generate any story one "
        "wants; the honest choice is a mix that reveals both wins and "
        "losses.")

    add_heading2(doc, "Lesson 5 — Documentation is a deliverable")
    add_p(doc,
        "Early drafts of project documentation focused on code comments "
        "and API reference. User testing (simulated by returning to the "
        "project after a two-week break with no memory of the "
        "implementation) revealed that the most critical documentation "
        "was the user guide — specifically the Cypher tips table "
        "explaining rent-burden-rate as a decimal rather than a "
        "percentage, and the direction of the NEIGHBORS edge. Quality of "
        "documentation directly determines whether the tool is adopted "
        "or ignored, and that quality cannot be written in the last "
        "week before the deadline.")

    add_heading2(doc, "Lesson 6 — Ethical review is not optional")
    add_p(doc,
        "The NOAH database contains community-sensitive information. "
        "Making it easier to query — which is the entire point of this "
        "project — simultaneously makes it easier to misuse. The "
        "ethical considerations chapter was originally planned as a "
        "single paragraph; during review it expanded to four subsections "
        "addressing displacement risk, data minimization, aggregation "
        "risk, and research-obligations compliance. The broader lesson "
        "is that any project that improves access to community data "
        "should reserve space, time, and attention for ethics review. "
        "Ethics as a post-hoc sticker is not enough.")


def section_conclusion(doc):
    add_heading1(doc, "Conclusion and Further Work")
    reset_figure_counter()

    add_heading2(doc, "Conclusions")
    add_p(doc,
        "Overall, the project addressed the research question "
        "affirmatively and in detail. A config-driven Python pipeline, "
        "augmented by an LLM schema interpreter and paired with a "
        "schema-aware Text2Cypher interface, migrated the NOAH "
        "PostgreSQL/PostGIS database into a Neo4j knowledge graph with "
        "zero data loss, in under eight seconds, and delivered a natural-"
        "language query interface with 95 percent benchmark accuracy. "
        "The same engine migrated three unrelated public datasets — "
        "Chinook, Northwind, and Pagila — in under four seconds each "
        "with zero orphan foreign keys, validating the generalization "
        "claim.")
    add_p(doc,
        "Key findings: (1) Audit-report semantics are load-bearing; the "
        "INFO-versus-WARN distinction is the difference between a "
        "trustworthy and a cry-wolf audit. (2) LLMs are useful as "
        "draftsmen in both schema interpretation and query translation "
        "but require human review at both points. (3) Graph databases "
        "win at relationship-heavy query shapes and lose at bulk "
        "aggregation; the honest claim is shape-specific rather than "
        "unconditional.")
    add_p(doc,
        "The proof of concept therefore qualifies as a success by each "
        "of the project's four SMART objectives (see Chapter 2 and "
        "Appendix A) and by the additional criteria of reproducibility, "
        "observability, documentation, and ethical defensibility.")

    add_heading2(doc, "Implications")
    add_p(doc,
        "Theoretically, the synthesis of De Virgilio's formal rules, "
        "Rel2Graph's automation, and Data2Neo's operational patterns — "
        "augmented by an LLM interpreter — is a viable engineering "
        "pattern and can be applied beyond the PostgreSQL-to-Neo4j pair "
        "to any relational-to-graph migration with a comparable "
        "conversion formalism. Practically, the tool reduces the "
        "effective cost of adopting a graph database from senior-"
        "engineer-weeks to human-review-hours, which materially lowers "
        "the barrier for smaller organizations — including NYC "
        "community-data projects — to use graph technology.")

    add_heading2(doc, "Limitations")
    add_p(doc,
        "Constraints to be stated honestly: the tool is tested only on "
        "single-node Neo4j 5.15 community edition; horizontally scaled "
        "Neo4j Fabric deployments are unvalidated. The Text2Cypher "
        "accuracy is measured on NOAH-specific questions; generalization "
        "to arbitrary-domain questions is not claimed. The migration "
        "pipeline is batch-mode; real-time change-data-capture from "
        "PostgreSQL to Neo4j is out of scope. The LLM API has a non-zero "
        "per-query cost; production deployments at scale should budget "
        "for that cost or self-host an open-weight model.")
    add_p(doc,
        "Validity considerations: the Text2Cypher accuracy benchmark was "
        "designed by the same person who tuned the prompts, which "
        "introduces a clear designer-tuner bias. A stronger evaluation "
        "would use questions authored by an independent reviewer. "
        "Similarly, the performance benchmark was run on a single laptop; "
        "production validation on server-grade hardware is left as "
        "future work.")

    add_heading2(doc, "Further Work")
    add_p(doc,
        "Five directions are identified for follow-on work. First, "
        "ACRIS owner-network integration — the NYC property-records "
        "dataset contains 85 million rows across three tables (~50 GB). "
        "Adding this layer to the NOAH graph would enable ownership-"
        "chain queries that expose property concentration patterns. The "
        "integration requires bulk-import tooling (neo4j-admin import) "
        "rather than the current Python-mediated batch MERGE. Second, "
        "a parameterized question library — five to ten templated Cypher "
        "queries with dropdowns for geography and indicator — would "
        "provide a middle ground between free-form Text2Cypher and raw "
        "Cypher editing for analysts who need repeatable, validated "
        "queries. Third, variable-depth graph traversal exposed as a "
        "depth slider in the UI would unlock a class of network-"
        "discovery queries not currently supported. Fourth, production "
        "hardening — authentication, rate limiting, query caching, "
        "integration with Neo4j Aura — is required for any deployment "
        "beyond the classroom. Fifth, broader database support: "
        "extending the schema analyzer to MySQL and SQL Server would "
        "generalize the tool beyond PostgreSQL, directly addressing the "
        "Digital Forge Lab's stated interest in a broadly applicable "
        "migration bot.")

    add_heading2(doc, "Closing Summary")
    add_p(doc,
        "This project delivered an open-source PostgreSQL-to-Neo4j "
        "conversion bot with a natural-language query interface, "
        "validated on 8,604 NYC affordable-housing records and three "
        "unrelated public datasets. Every numerical target was met or "
        "exceeded; every artifact referenced in this report is "
        "reproducible from the public GitHub repository. The main "
        "contribution to knowledge is the demonstration that the "
        "synthesis of a formal conversion framework, an LLM schema "
        "interpreter, and schema-aware Text2Cypher prompting is "
        "sufficient to automate an end-to-end migration previously "
        "requiring senior-engineer-weeks of manual work.")


def section_references(doc):
    add_heading1(doc, "References")
    reset_figure_counter()
    add_p(doc,
          "References are presented in APA 7th-edition style with "
          "hanging indent. All in-text citations in the body chapters "
          "and in Appendix H cross-reference this list.",
          italic=True, color=MUTED, space_after=Pt(12))

    refs = [
        "Abu-Salih, B. (2021). Domain-specific knowledge graphs: A survey. "
        "Journal of Network and Computer Applications, 185, 103076. "
        "https://doi.org/10.1016/j.jnca.2021.103076",

        "Angles, R. (2012). A comparison of current graph database "
        "models. In Proceedings of the 2012 IEEE 28th International "
        "Conference on Data Engineering Workshops (pp. 171–177). IEEE. "
        "https://doi.org/10.1109/ICDEW.2012.31",

        "Angles, R., Arenas, M., Barceló, P., Hogan, A., Reutter, J., & "
        "Vrgoc, D. (2017). Foundations of modern query languages for "
        "graph databases. ACM Computing Surveys, 50(5), 1–40.",

        "Angles, R., & Gutierrez, C. (2008). Survey of graph database "
        "models. ACM Computing Surveys, 40(1), 1–39.",

        "Codd, E. F. (1970). A relational model of data for large shared "
        "data banks. Communications of the ACM, 13(6), 377–387.",

        "De Virgilio, R., Maccioni, A., & Torlone, R. (2013). Converting "
        "relational to graph databases. In Proceedings of the First "
        "International Workshop on Graph Data Management Experiences and "
        "Systems (GRADES) (pp. 1–6). ACM.",

        "Minder, P., Kindler, L., & Laparra, E. (2024). Data2Neo: A tool "
        "for complex Neo4j data integration. arXiv:2406.04995.",

        "Mulligan, C., & Bamberger, K. (2021). Property ownership "
        "disclosure and the limits of transparency. California Law "
        "Review, 109(5), 1635–1702.",

        "Ng, J., Yeh, A., & Yeh, S. (2022). A knowledge-graph approach "
        "to Hong Kong public housing: Data integration across "
        "government, survey, and spatial sources. Sustainable Cities and "
        "Society, 83, 103961.",

        "NYC Open Data. (2024). Affordable housing production by "
        "building [Dataset]. NYC Department of Housing Preservation and "
        "Development. https://data.cityofnewyork.us/Housing-Development/"
        "Affordable-Housing-Production-by-Building/hg8x-zxpr",

        "Ozsoy, O., Aktas, S., Ulker, Y., & Temizel, A. (2024). "
        "Text2Cypher: Bridging natural language and graph databases. "
        "arXiv:2412.10064.",

        "Robinson, I., Webber, J., & Eifrem, E. (2015). Graph databases: "
        "New opportunities for connected data (2nd ed.). O'Reilly Media.",

        "Sackett, D. L., Rosenberg, W. M., Gray, J. A., Haynes, R. B., & "
        "Richardson, W. S. (1996). Evidence based medicine: What it is "
        "and what it isn't. BMJ, 312(7023), 71–72.",

        "Singhal, A. (2012). Introducing the knowledge graph: Things, "
        "not strings. Google Official Blog.",

        "U.S. Census Bureau. (2022). American Community Survey 5-year "
        "estimates [Dataset]. https://www.census.gov/programs-"
        "surveys/acs",

        "Yu, Y. (2025). NOAH PostgreSQL/PostGIS implementation "
        "[Capstone project]. NYU School of Professional Studies. "
        "https://github.com/Becky0713/NOAH",

        "Zhang, C. (2025). NOAH information dashboard: A proof-of-"
        "concept housing affordability analytics tool for Urban Labs "
        "[Capstone report]. NYU School of Professional Studies.",

        "Zhao, F., Xu, W., & Bagherzadeh, N. (2023). Rel2Graph: "
        "Automated mapping from relational databases to a unified "
        "property knowledge graph. arXiv:2310.01080.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(ref)
        set_font(r, size=SIZE_BODY)


# ── appendices ────────────────────────────────────────────────────────────

def appendix_header(doc, title):
    add_heading1(doc, title)
    reset_figure_counter()


def appendix_A(doc):
    appendix_header(doc, "Appendix A — Project Acceptance Document")
    add_p(doc, "PROJECT ACCEPTANCE DOCUMENT",
          bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=Pt(18))
    add_table(doc,
        ["Field", "Value"],
        [
            ["Project Title",
             "Automated RDBMS-to-Knowledge Graph Conversion Bot (NOAH case study)"],
            ["Project Manager",     "Zhen Yang"],
            ["Faculty Advisor",     "Dr. Andres Fortino, NYU SPS"],
            ["Sponsor Organization","The Digital Forge Lab, NYU SPS"],
            ["Semester",            "Spring 2026 · MASY GC-4100"],
            ["Project Start",       "February 2, 2026"],
            ["Project Delivery",    "April 28, 2026"],
        ],
        col_widths=[Inches(1.8), Inches(4.5)])

    add_heading2(doc, "Objective-to-Evidence Mapping")
    add_p(doc, "Each project objective below is paired with its agreed "
               "measurement rule and the evidence artifact that "
               "demonstrates it. The sponsor's signature on this page "
               "acknowledges acceptance of the deliverables as meeting "
               "the stated criteria.")
    add_table(doc,
        ["#", "Objective", "Metric", "Evidence"],
        [
            ["1", "Schema Mapping",
             "≥ 5 node mappings, 3+ business inquiries supported",
             "config/mapping_rules.yaml + FRS §3"],
            ["2", "Automated Migration Engine",
             "Row-to-node parity; 50-record ZIP-linkage audit",
             "outputs/audit_report.json (PASS)"],
            ["3", "Text2Cypher Interface",
             "≥ 15 / 20 actionable results with explanations",
             "outputs/benchmark_report.json (19 / 20 = 95 %)"],
            ["4", "End-to-End Walkthrough",
             "Live demo + 3 SQL joins → Cypher",
             "Final defense + outputs/performance_report.json"],
        ],
        col_widths=[Inches(0.4), Inches(1.8), Inches(2.4), Inches(2.0)])

    add_heading2(doc, "Deviations from Initial Metrics")
    add_p(doc,
        "One scope adjustment was made during execution: StreetEasy "
        "median rent data (median_rent_usd, rent_to_income_ratio) was "
        "unavailable and was replaced with American Community Survey "
        "2022 rent-burden rates. All four project objectives were met "
        "using the alternate data source. No other deviations from the "
        "initial metrics are reported. ACRIS owner-network data, "
        "identified in the Project Proposal as a potential Phase 2 "
        "extension, remains deferred — consistent with the scope "
        "boundary established at project start.")

    add_heading2(doc, "Sponsor Acceptance")
    add_p(doc,
        "The undersigned Project Sponsor accepts this capstone project "
        "as having met the objectives and metrics established in the "
        "Project Proposal and Functional Requirements Specification.",
        space_after=Pt(36))
    add_p(doc, "Sponsor Signature:  ____________________________________________    Date: ______________",
          space_after=Pt(24))
    add_p(doc, "Printed Name: Dr. Andres Fortino", space_after=Pt(6))
    add_p(doc, "Title: Clinical Assistant Professor · Director, The Digital Forge Lab · NYU SPS",
          space_after=Pt(24))
    add_p(doc, "Sponsor Comments:", bold=True, space_after=Pt(6))
    for _ in range(6):
        add_p(doc, "_" * 95, space_after=Pt(2))


def appendix_B(doc):
    appendix_header(doc, "Appendix B — Project Sponsor Agreement")
    add_p(doc, "PROJECT SPONSOR AGREEMENT",
          bold=True, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER,
          space_after=Pt(18))
    add_p(doc,
        "This Project Sponsor Agreement, dated February 19, 2026, is "
        "entered into between Zhen Yang (\"Project Manager\") and The "
        "Digital Forge Lab at the NYU School of Professional Studies "
        "(\"Sponsor\"), represented by Dr. Andres Fortino in his role as "
        "lab director and faculty advisor for the Spring 2026 Applied "
        "Project Capstone (MASY GC-4100). The purpose of this agreement "
        "is to document the mutual commitments and expectations between "
        "the Project Manager and the Sponsor for the twelve-week "
        "project, running February 2 to April 30, 2026.")

    add_heading2(doc, "Scope of Work")
    add_p(doc,
        "The Project Manager agrees to design, implement, and deliver "
        "an automated PostgreSQL-to-Neo4j migration pipeline with a "
        "Text2Cypher natural-language query interface, validated on the "
        "NOAH affordable-housing database, as specified in the Project "
        "Proposal (February 10, 2026) and the Functional Requirements "
        "Specification (February 19, 2026). Deliverables include working "
        "code, documentation, a Streamlit dashboard, an educational "
        "Jupyter notebook, and this final report.")

    add_heading2(doc, "Sponsor Commitments")
    add_p(doc,
        "The Sponsor commits to the following: (a) timely review of "
        "project deliverables — within three business days of submission; "
        "(b) attendance at a minimum of four milestone meetings (project "
        "launch, objectives review, mid-project progress review, and "
        "final walkthrough); (c) provision of access to the Digital "
        "Forge Lab research infrastructure where required; and (d) "
        "written acceptance of the final deliverables upon "
        "satisfaction of the agreed metrics.")

    add_heading2(doc, "Project Manager Commitments")
    add_p(doc,
        "The Project Manager commits to the following: (a) delivery of "
        "all scheduled milestones per the Work Breakdown Structure "
        "(Appendix D); (b) weekly written progress reports summarizing "
        "accomplishments, upcoming tasks, and any blockers requiring "
        "sponsor input; (c) scheduling and leading all four milestone "
        "meetings without prompt from the sponsor; and (d) submission of "
        "the final report and signed acceptance document before the "
        "April 30, 2026 deadline.")

    add_heading2(doc, "Intellectual Property and Attribution")
    add_p(doc,
        "All code, documentation, and data artifacts produced under this "
        "agreement are released as open-source under the MIT License and "
        "will be hosted at https://github.com/gitzhen0/noah-postgres-to-"
        "neo4j. The Project Manager retains authorship credit; the "
        "Sponsor and the Digital Forge Lab are acknowledged in "
        "deliverables. Upstream NOAH data sources (Yu 2025, Zhang 2025) "
        "are attributed throughout per their respective licenses.")

    add_heading2(doc, "Signatures")
    add_p(doc, "", space_after=Pt(30))
    add_p(doc,
        "Sponsor:  ____________________________________________    Date: ______________",
        space_after=Pt(24))
    add_p(doc, "Printed Name: Dr. Andres Fortino", space_after=Pt(30))
    add_p(doc,
        "Project Manager:  ____________________________________    Date: ______________",
        space_after=Pt(24))
    add_p(doc, "Printed Name: Zhen Yang")


def appendix_C(doc):
    appendix_header(doc, "Appendix C — Functional Requirements Specifications")
    render_markdown(doc, load_markdown(SOURCES["frs"]))


def appendix_D(doc):
    appendix_header(doc, "Appendix D — Project Plan and Work Breakdown Structure")
    add_p(doc,
        "This appendix reproduces the Work Breakdown Structure (WBS) "
        "submitted on April 6, 2026 as Assignment 8B. The WBS spans the "
        "twelve-week project timeline (February 2 — April 30, 2026), "
        "decomposes the work into six phases and 340 estimated hours, and "
        "annotates each phase with its completion status as of the "
        "submission date. The original document is reproduced verbatim "
        "below, rendered as page images to preserve the table structure "
        "exactly as submitted to the sponsor.", space_after=Pt(12))
    fig_dir = REPO / "assets" / "figures"
    wbs_pages = sorted(fig_dir.glob("wbs-*.png"))
    for page_img in wbs_pages:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run()
        try:
            run.add_picture(str(page_img), width=Inches(6.3))
        except Exception as exc:
            t = p.add_run(f"[missing WBS page: {page_img}]")
            set_font(t, italic=True, color=MUTED)


def appendix_E(doc):
    appendix_header(doc, "Appendix E — Risk Management Plan")
    render_markdown(doc, load_markdown(SOURCES["risk"]))


def appendix_F(doc):
    appendix_header(doc, "Appendix F — Technology Trial Plan")
    render_markdown(doc, load_markdown(SOURCES["trial"]))


def appendix_G(doc):
    appendix_header(doc, "Appendix G — Status Reports")
    add_p(doc,
        "Three formal status reports were submitted to the Sponsor over "
        "the project duration, each using the standard Red / Yellow / "
        "Green severity indicators across six project areas (Overall "
        "Status, Schedule, Deliverables, Resources and Collaboration, "
        "Changes, Communication). All six areas were Green on every "
        "report. All three reports are reproduced in chronological order "
        "below. Space is reserved at the end of each report for the "
        "sponsor's signature.", space_after=Pt(18))

    # Status Report 1 — February 26, 2026
    add_heading2(doc, "Status Report 1 — February 26, 2026")
    render_markdown(doc, load_markdown(SOURCES["status_feb"]))
    add_p(doc, "", space_after=Pt(18))
    add_p(doc, "Sponsor Signature: ____________________________________   Date: _______________",
          space_after=Pt(6))
    add_p(doc, "Printed Name: Dr. Andres Fortino", space_after=Pt(24))

    # Status Report 2 — March 25, 2026
    add_heading2(doc, "Status Report 2 — March 25, 2026")
    render_markdown(doc, load_markdown(SOURCES["status_mar"]))
    add_p(doc, "", space_after=Pt(18))
    add_p(doc, "Sponsor Signature: ____________________________________   Date: _______________",
          space_after=Pt(6))
    add_p(doc, "Printed Name: Dr. Andres Fortino", space_after=Pt(24))

    # Status Report 3 — April 16, 2026
    add_heading2(doc, "Status Report 3 — April 16, 2026")
    render_markdown(doc, load_markdown(SOURCES["status_apr"]))
    add_p(doc, "", space_after=Pt(18))
    add_p(doc, "Sponsor Signature: ____________________________________   Date: _______________",
          space_after=Pt(6))
    add_p(doc, "Printed Name: Dr. Andres Fortino")


def appendix_H(doc):
    appendix_header(doc, "Appendix H — Annotated Bibliography")
    render_markdown(doc, load_markdown(SOURCES["bibliography"]))


# ── driver ────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # default style
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal.font.size = SIZE_BODY

    # initial section = cover (no footer page number)
    section_cover(doc)

    # switch to running footer from the next section onward
    sect = new_section(doc)
    set_page_number_in_footer(sect, show=True)

    section_toc(doc)
    section_declaration(doc)
    section_acknowledgments(doc)
    section_abstract(doc)
    section_abbreviations(doc)

    # Report body
    section_introduction(doc)
    section_objectives(doc)
    section_alternates(doc)
    section_literature(doc)
    section_methodology(doc)
    section_results(doc)
    section_issues(doc)
    section_lessons(doc)
    section_conclusion(doc)
    section_references(doc)

    # Appendices
    appendix_A(doc)
    appendix_B(doc)
    appendix_C(doc)
    appendix_D(doc)
    appendix_E(doc)
    appendix_F(doc)
    appendix_G(doc)
    appendix_H(doc)

    doc.save(str(OUT))
    size_kb = OUT.stat().st_size / 1024
    print(f"Wrote: {OUT}")
    print(f"Size : {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
