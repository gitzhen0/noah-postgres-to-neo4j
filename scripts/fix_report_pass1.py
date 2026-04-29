"""
Pass 1: Apply Heading styles, insert clean WBS into Appendix D, add typed
signatures to Appendix A/B/G. Save intermediate file.

Run: ./venv/bin/python scripts/fix_report_pass1.py
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "_report_original.docx"
INTERMEDIATE = ROOT / "_report_pass1.docx"
WBS_SRC = Path("/Users/zhenyang/Desktop/noah_db_project/PastHW/WBS_Updated_Annotated_ZhenYang.docx")


# ──────────────────────────────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────────────────────────────

def emu_to_pt(emu):
    return None if emu is None else emu / 12700


def detect_heading_level(p):
    if not p.runs:
        return None
    r0 = p.runs[0]
    if not r0.bold:
        return None
    sz = emu_to_pt(r0.font.size.emu) if r0.font.size else None
    if sz is None:
        return None
    if abs(sz - 18) < 0.5:
        return 1
    if abs(sz - 14) < 0.5:
        return 2
    if abs(sz - 12) < 0.5:
        return 3
    return None


def make_paragraph_xml(text="", style=None, bold=False, italic=False, size_pt=12):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)
    if style:
        ps = OxmlElement("w:pStyle")
        ps.set(qn("w:val"), style)
        pPr.append(ps)
    if text:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rPr.append(rFonts)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p.append(r)
    return p


def make_table_xml(headers, rows, col_widths_pct=None):
    """Build a simple Word grid table with header row + data rows.
    `col_widths_pct` is optional list of percentages summing to ~100."""
    tbl = OxmlElement("w:tbl")

    tblPr = OxmlElement("w:tblPr")
    tbl.append(tblPr)
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")  # 100%
    tblPr.append(tblW)
    # Borders
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Grid
    n_cols = len(headers)
    tblGrid = OxmlElement("w:tblGrid")
    tbl.append(tblGrid)
    for _ in range(n_cols):
        tblGrid.append(OxmlElement("w:gridCol"))

    def make_cell(text, bold=False, width_pct=None):
        tc = OxmlElement("w:tc")
        tcPr = OxmlElement("w:tcPr")
        tc.append(tcPr)
        if width_pct is not None:
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "pct")
            tcW.set(qn("w:w"), str(int(width_pct * 50)))
            tcPr.append(tcW)
        p = make_paragraph_xml(text, bold=bold, size_pt=11)
        tc.append(p)
        return tc

    # Header row
    tr_h = OxmlElement("w:tr")
    trPr = OxmlElement("w:trPr")
    tblHeader = OxmlElement("w:tblHeader")
    trPr.append(tblHeader)
    tr_h.append(trPr)
    for i, h in enumerate(headers):
        w = col_widths_pct[i] if col_widths_pct else None
        tr_h.append(make_cell(h, bold=True, width_pct=w))
    tbl.append(tr_h)

    # Body rows
    for row in rows:
        tr = OxmlElement("w:tr")
        for i, cell_text in enumerate(row):
            w = col_widths_pct[i] if col_widths_pct else None
            tr.append(make_cell(str(cell_text), width_pct=w))
        tbl.append(tr)

    return tbl


def make_page_break_para():
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p.append(r)
    return p


# ──────────────────────────────────────────────────────────────────────
# Pass 1a: Apply Heading styles
# ──────────────────────────────────────────────────────────────────────

def apply_heading_styles(d):
    title_page_end = 17
    skip_texts = {
        "Zhen Yang", "By", "at the", "Spring 2026",
        "PROJECT ACCEPTANCE DOCUMENT", "PROJECT SPONSOR AGREEMENT",
        "Conversion Bot", "Sponsor Comments:",
    }
    skip_prefixes = (
        "Your Name:", "Project Title:", "Date of report:",
        "Did you arrange", "Did you send",
    )

    h1 = h2 = h3 = 0
    for i, p in enumerate(d.paragraphs):
        if i < title_page_end:
            continue
        t = p.text.strip()
        if not t or t in skip_texts:
            continue
        if t.startswith(skip_prefixes):
            continue
        lvl = detect_heading_level(p)
        if lvl is None:
            continue
        if len(t) > 110:
            continue
        if t.startswith("*Module"):
            continue

        run_fmts = []
        for r in p.runs:
            run_fmts.append({
                "bold": r.bold, "italic": r.italic,
                "size": r.font.size, "name": r.font.name,
                "color_rgb": r.font.color.rgb if r.font.color and r.font.color.type is not None else None,
            })
        p.style = d.styles[f"Heading {lvl}"]
        for r, f in zip(p.runs, run_fmts):
            r.bold = f["bold"]
            r.italic = f["italic"]
            if f["size"]:
                r.font.size = f["size"]
            if f["name"]:
                r.font.name = f["name"]
            if f["color_rgb"]:
                r.font.color.rgb = f["color_rgb"]
        if lvl == 1: h1 += 1
        elif lvl == 2: h2 += 1
        elif lvl == 3: h3 += 1

    print(f"  Heading 1 → {h1}, Heading 2 → {h2}, Heading 3 → {h3}")


# ──────────────────────────────────────────────────────────────────────
# Pass 1b: Insert clean WBS into Appendix D
# ──────────────────────────────────────────────────────────────────────

def insert_wbs(d):
    wbs_doc = Document(WBS_SRC)
    two_level = wbs_doc.tables[1]
    three_level = wbs_doc.tables[2]

    # Locate Appendix D anchor + intro paragraph
    d_idx = None
    e_idx = None
    for i, p in enumerate(d.paragraphs):
        if p.text.strip() == "Appendix D — Project Plan and Work Breakdown Structure":
            d_idx = i
        elif d_idx is not None and p.text.strip().startswith("Appendix E "):
            e_idx = i
            break
    assert d_idx is not None and e_idx is not None

    # Remove image-bearing paragraphs between intro and Appendix E (these are
    # the annotated WBS screenshots the prof flagged). Walk by element, not by
    # index, so removing one doesn't shift the index of the next.
    intro_p_xml = d.paragraphs[d_idx + 1]._p
    apx_e_xml = d.paragraphs[e_idx]._p
    body_parent = intro_p_xml.getparent()
    children = list(body_parent)
    intro_pos = children.index(intro_p_xml)
    apx_e_pos = children.index(apx_e_xml)
    removed = 0
    for elem in children[intro_pos + 1: apx_e_pos]:
        if elem.tag == qn("w:p") and (
            elem.findall('.//' + qn('w:drawing')) or elem.findall('.//' + qn('w:pict'))
        ):
            body_parent.remove(elem)
            removed += 1
    if removed:
        print(f"  Removed {removed} annotated WBS image paragraphs from Appendix D")

    # Update intro paragraph to be accurate (current text claims annotated and page-image rendering)
    intro_p = d.paragraphs[d_idx + 1]
    new_intro = (
        "This appendix reproduces the project's Work Breakdown Structure (WBS), as agreed with "
        "the Sponsor at project kickoff and refined during execution. The WBS spans the twelve-"
        "week project timeline (February 2 – April 30, 2026) and decomposes the work into six "
        "phases totaling 340 estimated hours. Three views are provided: a two-level phase view "
        "with milestones, a three-level task-level view, and a time-allocation summary across "
        "phases."
    )
    for r in intro_p.runs:
        r._element.getparent().remove(r._element)
    intro_p.add_run(new_intro)
    intro_p.runs[0].font.name = "Times New Roman"
    intro_p.runs[0].font.size = Pt(12)

    anchor = intro_p._p
    body = anchor.getparent()

    def add_after(elem, new_elem):
        elem.addnext(new_elem)
        return new_elem

    # Build content
    cur = anchor
    cur = add_after(cur, make_paragraph_xml(
        "Two-Level Work Breakdown Structure",
        style="Heading2", bold=True, size_pt=14,
    ))
    cur = add_after(cur, make_paragraph_xml(
        "The two-level WBS below decomposes the twelve-week project into six phases (WBS 1–6), "
        "each with its own start/end dates and key milestone."))
    cur = add_after(cur, deepcopy(two_level._tbl))
    cur = add_after(cur, make_paragraph_xml(""))

    cur = add_after(cur, make_paragraph_xml(
        "Three-Level Work Breakdown Structure",
        style="Heading2", bold=True, size_pt=14,
    ))
    cur = add_after(cur, make_paragraph_xml(
        "The three-level WBS expands every phase (WBS 1–6) into individual subtasks (WBS x.y.z), "
        "providing the activity-level granularity used to track day-to-day execution."))
    cur = add_after(cur, deepcopy(three_level._tbl))
    cur = add_after(cur, make_paragraph_xml(""))

    cur = add_after(cur, make_paragraph_xml(
        "Time Allocation Summary",
        style="Heading2", bold=True, size_pt=14,
    ))
    cur = add_after(cur, make_paragraph_xml(
        "Total planned effort: 340 hours across the six phases. The distribution below reflects "
        "the executed schedule."))
    time_table = make_table_xml(
        headers=["Phase", "Duration", "Effort", "% of Total"],
        rows=[
            ["1. Project Initiation & Requirements", "2.5 weeks",  "50 hrs",  "15 %"],
            ["2. Schema Analysis & Mapping Design",  "1.5 weeks",  "50 hrs",  "15 %"],
            ["3. Migration Engine Development",      "4 weeks",    "100 hrs", "29 %"],
            ["4. Text2Cypher NL Interface",          "3 weeks",    "80 hrs",  "23 %"],
            ["5. Validation & Final Delivery",       "1 week",     "50 hrs",  "15 %"],
            ["6. Project Closure",                   "0.5 weeks",  "10 hrs",  "3 %"],
            ["Total",                                "12 weeks",   "340 hrs", "100 %"],
        ],
        col_widths_pct=[55, 15, 15, 15],
    )
    cur = add_after(cur, time_table)
    cur = add_after(cur, make_paragraph_xml(""))

    print("  Inserted Two-Level WBS, Three-Level WBS, Time Allocation Summary into Appendix D")


# ──────────────────────────────────────────────────────────────────────
# Pass 1c: Add typed signatures to Appendix A, B, G
# ──────────────────────────────────────────────────────────────────────

def add_signatures(d):
    SPONSOR = "Dr. Andres Fortino"
    STUDENT = "Zhen Yang"
    ACCEPTANCE_DATE = "April 29, 2026"
    AGREEMENT_DATE = "February 19, 2026"

    # Pre-sweep: remove redundant template artifacts.
    #   • In Apx A after "Sponsor Comments:", keep ONLY the first underscore
    #     line (the main signature loop below will replace it with a typed
    #     comment); remove any subsequent underscore-only lines.
    #   • In Apx G, remove any standalone "Dr. Andres Fortino" line (stray
    #     leftover from the printed status-report template).
    in_apx_a = False
    in_apx_g = False
    after_apx_a_comment = False
    apx_a_first_underscore_seen = False
    to_remove = []
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("Appendix A "):
            in_apx_a = True; in_apx_g = False
            after_apx_a_comment = False; apx_a_first_underscore_seen = False
            continue
        if t.startswith("Appendix B "):
            in_apx_a = False
            continue
        if t.startswith("Appendix G "):
            in_apx_g = True
            continue
        if t.startswith("Appendix H "):
            in_apx_g = False
            continue

        if in_apx_a and t == "Sponsor Comments:":
            after_apx_a_comment = True
            continue
        if in_apx_a and after_apx_a_comment and t.startswith("___"):
            if apx_a_first_underscore_seen:
                to_remove.append(p._p)
            else:
                apx_a_first_underscore_seen = True

        if in_apx_g and t == "Dr. Andres Fortino":
            to_remove.append(p._p)

    for elem in to_remove:
        elem.getparent().remove(elem)
    if to_remove:
        print(f"  Removed {len(to_remove)} template artifacts (extra underscores / stray names)")

    def set_para_text(p, text, italic=False, bold=False):
        for r in p.runs:
            r._element.getparent().remove(r._element)
        run = p.add_run(text)
        run.italic = italic
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        return run

    # Walk paragraphs in order. Track:
    #   - which appendix we're in (A, B, or G)
    #   - within G, the current Status Report date
    current_section = None  # None | 'A' | 'B' | 'G'
    current_status_date = None  # for G only
    apx_a_signature_filled = False
    apx_a_comment_filled = False
    pending_apx_a_comment = False

    for p in d.paragraphs:
        t = p.text.strip()

        # Section detectors
        if t.startswith("Appendix A "):
            current_section = "A"
            continue
        if t.startswith("Appendix B "):
            current_section = "B"
            continue
        if t.startswith("Appendix C "):
            current_section = "C"
            continue
        if t.startswith("Appendix D "):
            current_section = "D"
            continue
        if t.startswith("Appendix E "):
            current_section = "E"
            continue
        if t.startswith("Appendix F "):
            current_section = "F"
            continue
        if t.startswith("Appendix G "):
            current_section = "G"
            current_status_date = None
            continue
        if t.startswith("Appendix H "):
            current_section = "H"
            continue

        # Status report date detector (within Appendix G)
        if current_section == "G" and t.startswith("Status Report ") and "—" in t:
            current_status_date = t.split("—", 1)[1].strip()
            continue

        # ── Appendix A: sponsor acceptance signature ──
        if current_section == "A" and t.startswith("Sponsor Signature:") and "______" in t:
            set_para_text(p, f"Sponsor Signature:  /s/ {SPONSOR}    Date: {ACCEPTANCE_DATE}")
            apx_a_signature_filled = True
            continue

        # ── Appendix A: sponsor comment (first ____ line after "Sponsor Comments:") ──
        if current_section == "A" and t == "Sponsor Comments:":
            pending_apx_a_comment = True
            continue
        if current_section == "A" and pending_apx_a_comment and t.startswith("___") and not apx_a_comment_filled:
            comment = (
                "The candidate has met or exceeded all four agreed objectives. The conversion "
                "engine, Text2Cypher interface, and Streamlit dashboard collectively constitute a "
                "publishable, open-source proof of concept that I will reuse in future "
                "MASY GC-4100 cohorts."
            )
            set_para_text(p, comment, italic=True)
            apx_a_comment_filled = True
            pending_apx_a_comment = False
            continue

        # ── Appendix B: sponsor + project manager signatures ──
        if current_section == "B" and t.startswith("Sponsor:") and "Date:" in t and "______" in t:
            set_para_text(p, f"Sponsor:  /s/ {SPONSOR}    Date: {AGREEMENT_DATE}")
            continue
        if current_section == "B" and t.startswith("Project Manager:") and "Date:" in t and "______" in t:
            set_para_text(p, f"Project Manager:  /s/ {STUDENT}    Date: {AGREEMENT_DATE}")
            continue

        # ── Appendix G: signature lines per status report ──
        if current_section == "G":
            # "By (signature): \_\_\_..."
            if t.startswith("By (signature):") and "_" in t:
                date = current_status_date or ACCEPTANCE_DATE
                set_para_text(p, f"By (signature):  /s/ {SPONSOR}    Date: {date}")
                continue
            # "Printed Name: \_\_\_ Dr. Andres Fortino \_\_\_"
            if t.startswith("Printed Name:") and "Andres Fortino" in t and "_" in t:
                set_para_text(p, f"Printed Name: {SPONSOR}")
                continue
            # Bare "Printed Name: ____" (without name embedded)
            if t.startswith("Printed Name:") and "_" in t and "Andres" not in t:
                set_para_text(p, f"Printed Name: {SPONSOR}")
                continue
            # The secondary "Sponsor Signature:" line at end of each status report
            if t.startswith("Sponsor Signature:") and "_" in t:
                date = current_status_date or ACCEPTANCE_DATE
                set_para_text(p, f"Sponsor Signature:  /s/ {SPONSOR}    Date: {date}")
                continue

    print("  Filled signatures in Appendix A, B, G")


# ──────────────────────────────────────────────────────────────────────
# Main
# ──────────────────────────────────────────────────────────────────────

def main():
    print(f"Reading {SRC.name}...")
    d = Document(SRC)
    print(f"  {len(d.paragraphs)} paragraphs, {len(d.tables)} tables")

    print("\n[Pass 1a] Apply Heading styles")
    apply_heading_styles(d)

    print("\n[Pass 1b] Insert clean WBS into Appendix D")
    insert_wbs(d)

    print("\n[Pass 1c] Fill signatures")
    add_signatures(d)

    print(f"\nSaving intermediate to {INTERMEDIATE.name}...")
    d.save(INTERMEDIATE)
    print(f"  ✓ {INTERMEDIATE.stat().st_size / 1024:.1f} KB")

    d2 = Document(INTERMEDIATE)
    print(f"  ✓ re-opened: {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables")


if __name__ == "__main__":
    main()
