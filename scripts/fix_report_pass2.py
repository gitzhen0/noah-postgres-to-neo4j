"""
Pass 2: Insert real Table of Contents + List of Figures with computed page
numbers. Iterative — detects page shift caused by TOC/LOF growth and re-runs.

Run: ./venv/bin/python scripts/fix_report_pass2.py
"""

from copy import deepcopy
from pathlib import Path
import json
import re
import subprocess

import fitz
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "_report_pass1.docx"
OUT = ROOT / "NOAH_Capstone_Final_Report_fixed.docx"
TMP_PDF = ROOT / "_report_final_check.pdf"

H1_TITLES = [
    "Declaration",
    "Acknowledgments",
    "Abstract",
    "Abbreviations and Definitions",
    "Introduction",
    "Project Objectives and Metrics",
    "Alternate Solutions Evaluated",
    "Literature Survey",
    "Approach and Methodology",
    "Results",
    "Issues Encountered",
    "Lessons Learned",
    "Conclusion and Further Work",
    "References",
    "Appendix A — Project Acceptance Document",
    "Appendix B — Project Sponsor Agreement",
    "Appendix C — Functional Requirements Specifications",
    "Appendix D — Project Plan and Work Breakdown Structure",
    "Appendix E — Risk Management Plan",
    "Appendix F — Technology Trial Plan",
    "Appendix G — Status Reports",
    "Appendix H — Annotated Bibliography",
]

# Curated H2 subsections to include under their parent H1, in document order.
H2_BY_PARENT = {
    "Introduction": [
        "Problem", "Approach", "Core Technology", "Benefits",
        "Research Question", "Contribution", "Sponsor", "Importance of Project",
    ],
    "Project Objectives and Metrics": [
        "Goal of the project", "Project Deliverables and Metrics", "Project Evaluation",
    ],
    "Alternate Solutions Evaluated": [
        "Solution A: Hand-written ETL Script Per Dataset",
        "Solution B: Off-the-Shelf Enterprise ETL (Informatica, Talend)",
        "Solution C: Config-Driven Pipeline with LLM Augmentation",
        "Solution Evaluation Criteria", "Selection Rationale",
    ],
    "Literature Survey": [
        "The Industry", "The Problem", "The Proposed Solution",
        "The Technology", "Use Cases",
    ],
    "Approach and Methodology": [
        "Problem Statement and Research Question", "Proof of Concept Approach",
        "Technology Trial Plan", "Population and Data", "Procedures",
        "Data Collection Methodology", "Data Analysis", "Organizational Change Plan",
    ],
    "Results": [
        "Data Processing", "Findings", "Summary Statistics",
        "Qualitative Observations", "Outcomes", "Implications", "Summary",
        "Repository of Data Sets and Code",
    ],
    "Issues Encountered": [
        "Risk Management Plan",
    ],
    "Lessons Learned": [],
    "Conclusion and Further Work": [
        "Conclusions", "Implications", "Limitations", "Further Work",
        "Closing Summary",
    ],
}

FIGURE_FULL_TEXT = {
    "Figure 5-1": "Figure 5-1: System architecture",
    "Figure 5-2": "Figure 5-2: Target NOAH graph model",
    "Figure 6-1": "Figure 6-1: Performance by query category",
    "Figure 6-2": "Figure 6-2: The hero query",
}
FIGURE_LABELS = {
    "Figure 5-1": "Figure 5-1 — System architecture",
    "Figure 5-2": "Figure 5-2 — Target NOAH graph model",
    "Figure 6-1": "Figure 6-1 — Performance by query category (PostgreSQL vs Neo4j)",
    "Figure 6-2": "Figure 6-2 — Hero query (2-hop ZIP traversal): Neo4j 37.7× faster",
}


# ──────────────────────────────────────────────────────────────────────
# Page extraction helpers
# ──────────────────────────────────────────────────────────────────────

def render_pdf(docx_path: Path, pdf_path: Path):
    if pdf_path.exists():
        pdf_path.unlink()
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf",
         "--outdir", str(pdf_path.parent), str(docx_path)],
        check=True, capture_output=True,
    )
    # soffice writes with same stem as input
    auto_pdf = pdf_path.parent / f"{docx_path.stem}.pdf"
    if auto_pdf != pdf_path:
        auto_pdf.rename(pdf_path)


def extract_pages(pdf_path: Path):
    """Extract page numbers for headings/figures. Auto-skips TOC pages by
    locating where Declaration body text starts ('I, Zhen Yang, declare')."""
    doc = fitz.open(pdf_path)

    def normalize(s): return re.sub(r"\s+", " ", s).strip()

    # Find first page where Declaration body actually starts. Anything earlier
    # is title-page or TOC; ignore for searches.
    DECL_BODY_MARKER = "I, Zhen Yang, declare"
    content_start_idx = 0  # 0-based
    for i in range(doc.page_count):
        if DECL_BODY_MARKER in doc.load_page(i).get_text():
            content_start_idx = i
            break

    def page_top_text(i):
        text = doc.load_page(i).get_text()
        norm = normalize(text)
        m = re.match(r"^(\d{1,4})\s+", norm)
        if m:
            norm = norm[m.end():]
        return norm[:400].lower()

    def page_full_text_lower(i):
        return normalize(doc.load_page(i).get_text()).lower()

    def find_h1(needle, after=content_start_idx):
        ndl = normalize(needle).lower()
        ndl_short = ndl.split(" — ")[0]
        first_words = " ".join(ndl.split()[:4])
        for i in range(after, doc.page_count):
            head = page_top_text(i)
            if ndl in head:
                return i + 1
            if len(ndl) > 25 and first_words in head and len(first_words) > 12:
                return i + 1
            if ndl_short and ndl_short != ndl and ndl_short in head:
                return i + 1
        return None

    h1_pages = {}
    last = content_start_idx
    for h in H1_TITLES:
        pg = find_h1(h, after=last)
        h1_pages[h] = pg
        if pg:
            last = pg - 1

    # H2 keyed by (parent_h1, h2_name) to avoid collisions when the same
    # H2 name (e.g. "Implications") appears under multiple chapters.
    h1_order = list(H1_TITLES)
    h2_pages = {}
    for parent_idx, parent_h1 in enumerate(h1_order):
        h2_children = H2_BY_PARENT.get(parent_h1, [])
        if not h2_children:
            continue
        parent_page = h1_pages.get(parent_h1)
        if not parent_page:
            for h2 in h2_children:
                h2_pages[(parent_h1, h2)] = None
            continue
        next_page = None
        for nxt in h1_order[parent_idx + 1:]:
            np = h1_pages.get(nxt)
            if np:
                next_page = np
                break
        upper_idx = (next_page - 1) if next_page else doc.page_count

        for h2 in h2_children:
            ndl = normalize(h2).lower()
            found = None
            for i in range(parent_page - 1, upper_idx):
                if ndl in page_full_text_lower(i):
                    found = i + 1
                    break
            h2_pages[(parent_h1, h2)] = found

    fig_pages = {}
    for short, full in FIGURE_FULL_TEXT.items():
        for i in range(content_start_idx, doc.page_count):
            if full in doc.load_page(i).get_text():
                fig_pages[short] = i + 1
                break
        else:
            fig_pages[short] = None

    return {
        "h1": h1_pages,
        "h2": h2_pages,
        "fig": fig_pages,
        "total": doc.page_count,
        "content_start_page": content_start_idx + 1,
    }


# ──────────────────────────────────────────────────────────────────────
# TOC building
# ──────────────────────────────────────────────────────────────────────

def make_toc_paragraph(text, page_num, level=1, doc=None):
    """Build a TOC entry: text [tab]....[tab] page_num.

    Uses tab stops with dot leader to right-aligned page number.
    """
    p = doc.paragraphs[0]._p.makeelement(qn("w:p"), {})
    # pPr with tab stops
    pPr = p.makeelement(qn("w:pPr"), {})
    p.append(pPr)
    # Indent for level 2
    if level == 2:
        ind = pPr.makeelement(qn("w:ind"), {qn("w:left"): "360"})
        pPr.append(ind)
    # Tabs
    tabs = pPr.makeelement(qn("w:tabs"), {})
    pPr.append(tabs)
    tab = tabs.makeelement(qn("w:tab"), {
        qn("w:val"): "right",
        qn("w:leader"): "dot",
        qn("w:pos"): "9000",
    })
    tabs.append(tab)

    # Run for entry text
    r1 = p.makeelement(qn("w:r"), {})
    p.append(r1)
    rPr1 = r1.makeelement(qn("w:rPr"), {})
    r1.append(rPr1)
    rFonts = rPr1.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman", qn("w:hAnsi"): "Times New Roman",
    })
    rPr1.append(rFonts)
    sz = rPr1.makeelement(qn("w:sz"), {qn("w:val"): "24"})  # 12pt
    rPr1.append(sz)
    if level == 1:
        b = rPr1.makeelement(qn("w:b"), {})
        rPr1.append(b)
    t1 = r1.makeelement(qn("w:t"), {})
    t1.text = text
    t1.set(qn("xml:space"), "preserve")
    r1.append(t1)

    # Tab + page number
    r2 = p.makeelement(qn("w:r"), {})
    p.append(r2)
    rPr2 = r2.makeelement(qn("w:rPr"), {})
    r2.append(rPr2)
    rFonts2 = rPr2.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman", qn("w:hAnsi"): "Times New Roman",
    })
    rPr2.append(rFonts2)
    sz2 = rPr2.makeelement(qn("w:sz"), {qn("w:val"): "24"})
    rPr2.append(sz2)
    tab_elem = r2.makeelement(qn("w:tab"), {})
    r2.append(tab_elem)
    t2 = r2.makeelement(qn("w:t"), {})
    t2.text = str(page_num) if page_num is not None else "—"
    r2.append(t2)

    return p


def make_section_heading(text, doc):
    """A bold 14pt centered heading paragraph."""
    p = doc.paragraphs[0]._p.makeelement(qn("w:p"), {})
    pPr = p.makeelement(qn("w:pPr"), {})
    p.append(pPr)
    pStyle = pPr.makeelement(qn("w:pStyle"), {qn("w:val"): "TOCHeading"})
    pPr.append(pStyle)
    jc = pPr.makeelement(qn("w:jc"), {qn("w:val"): "center"})
    pPr.append(jc)
    spc = pPr.makeelement(qn("w:spacing"), {qn("w:before"): "240", qn("w:after"): "120"})
    pPr.append(spc)

    r = p.makeelement(qn("w:r"), {})
    p.append(r)
    rPr = r.makeelement(qn("w:rPr"), {})
    r.append(rPr)
    rFonts = rPr.makeelement(qn("w:rFonts"), {
        qn("w:ascii"): "Times New Roman", qn("w:hAnsi"): "Times New Roman",
    })
    rPr.append(rFonts)
    b = rPr.makeelement(qn("w:b"), {})
    rPr.append(b)
    sz = rPr.makeelement(qn("w:sz"), {qn("w:val"): "32"})  # 16pt
    rPr.append(sz)
    color = rPr.makeelement(qn("w:color"), {qn("w:val"): "000000"})
    rPr.append(color)
    t = r.makeelement(qn("w:t"), {})
    t.text = text
    r.append(t)
    return p


def make_blank_para(doc):
    p = doc.paragraphs[0]._p.makeelement(qn("w:p"), {})
    return p


def build_toc_lof_elements(doc, pages):
    """Build a list of XML <w:p> elements: ToC then LoF."""
    elements = []
    # No "Table of Contents" heading here — paragraph 19 ("Table of Contents")
    # already exists; we only replace the placeholder text rows that follow.

    # H1 + H2 entries in document order. H2 dict is keyed (parent_h1, h2_name).
    parent_lookup = {h1: h2list for h1, h2list in H2_BY_PARENT.items()}
    for h1 in H1_TITLES:
        elements.append(make_toc_paragraph(h1, pages["h1"].get(h1), level=1, doc=doc))
        for h2 in parent_lookup.get(h1, []):
            pg = pages["h2"].get((h1, h2))
            elements.append(make_toc_paragraph(h2, pg, level=2, doc=doc))

    # List of Figures
    elements.append(make_blank_para(doc))
    elements.append(make_section_heading("List of Figures", doc))
    for short, label in FIGURE_LABELS.items():
        elements.append(make_toc_paragraph(label, pages["fig"].get(short), level=1, doc=doc))

    # Page break after LOF so following content (Declaration) starts fresh
    pb = doc.paragraphs[0]._p.makeelement(qn("w:p"), {})
    pPr = pb.makeelement(qn("w:pPr"), {})
    pb.append(pPr)
    r = pb.makeelement(qn("w:r"), {})
    pb.append(r)
    br = r.makeelement(qn("w:br"), {qn("w:type"): "page"})
    r.append(br)
    elements.append(pb)

    return elements


def insert_toc(doc, pages):
    """Replace placeholder paragraphs 20-21 with real TOC + LOF."""
    # Locate paragraph 19 ("Table of Contents") and the placeholder text after.
    # Strategy: find "Table of Contents" heading, then locate the placeholder text
    # and the existing "Declaration" heading. Anything between is replaced.
    body = doc.element.body
    paras = doc.paragraphs

    toc_heading_idx = None
    declaration_idx = None
    for i, p in enumerate(paras):
        if p.text.strip() == "Table of Contents" and toc_heading_idx is None:
            toc_heading_idx = i
        if p.text.strip() == "Declaration" and declaration_idx is None:
            declaration_idx = i
            break

    assert toc_heading_idx is not None and declaration_idx is not None

    # Remove paragraphs (toc_heading_idx + 1) .. (declaration_idx - 1) inclusive
    # Then insert TOC+LOF elements before declaration paragraph
    placeholder_paras = paras[toc_heading_idx + 1: declaration_idx]
    # Remove them from XML
    for pp in placeholder_paras:
        pp._p.getparent().remove(pp._p)

    # Re-fetch paragraphs to get fresh declaration anchor (XML now changed)
    paras = doc.paragraphs
    decl_p = None
    for p in paras:
        if p.text.strip() == "Declaration":
            decl_p = p
            break
    decl_anchor = decl_p._p

    # Insert TOC + LOF elements right before decl_anchor
    new_elems = build_toc_lof_elements(doc, pages)
    parent = decl_anchor.getparent()
    decl_idx_in_parent = list(parent).index(decl_anchor)
    for offset, elem in enumerate(new_elems):
        parent.insert(decl_idx_in_parent + offset, elem)


# ──────────────────────────────────────────────────────────────────────
# Iteration loop: insert TOC, render, check pages, redo if shifted
# ──────────────────────────────────────────────────────────────────────

def main():
    print("Iteration 1: render intermediate to get baseline page numbers")
    intermediate_pdf = ROOT / "_report_pass1.pdf"
    if not intermediate_pdf.exists():
        render_pdf(SRC, intermediate_pdf)
    pages = extract_pages(intermediate_pdf)
    print(f"  Total pages: {pages['total']}")

    # Iterate: insert TOC with current `pages`, render, re-extract `actual`, compare
    for iteration in range(1, 6):
        print(f"\n=== Iteration {iteration}: build TOC with current page numbers ===")
        d = Document(SRC)
        insert_toc(d, pages)
        d.save(OUT)
        print(f"  Saved candidate to {OUT.name}")

        render_pdf(OUT, TMP_PDF)
        new_pages = extract_pages(TMP_PDF)
        print(f"  Rendered: total={new_pages['total']} pages, "
              f"content starts at PDF page {new_pages['content_start_page']}")

        mismatches = []
        for h, expected in pages["h1"].items():
            actual = new_pages["h1"].get(h)
            if expected != actual:
                mismatches.append(("H1", h, expected, actual))
        for key, expected in pages["h2"].items():
            actual = new_pages["h2"].get(key)
            if expected != actual:
                mismatches.append(("H2", f"{key[0]}>{key[1]}", expected, actual))
        for f, expected in pages["fig"].items():
            actual = new_pages["fig"].get(f)
            if expected != actual:
                mismatches.append(("FIG", f, expected, actual))

        if not mismatches:
            print("  ✓ All page numbers match — TOC is converged.")
            break

        print(f"  ✗ {len(mismatches)} page mismatches; re-running with corrected pages.")
        for kind, h, e, a in mismatches[:10]:
            print(f"     [{kind}] {h!r:55} expected={e} actual={a}")
        pages = new_pages
    else:
        print("⚠ Did not converge in 5 iterations — leaving last attempt.")

    print(f"\n✓ Final report: {OUT}")
    print(f"  size: {OUT.stat().st_size / 1024:.1f} KB")
    print(f"  pages: {new_pages['total']}")


if __name__ == "__main__":
    main()
